"""Generate the revised MJDYPV submission package for the KFD manuscript."""

from __future__ import annotations

import re
from datetime import datetime
from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt


BASE_DIR = Path(__file__).resolve().parent.parent
MANUSCRIPT_DIR = BASE_DIR / "manuscripts"
REV_TABLES = BASE_DIR / "outputs" / "revision_tables"
REV_FIGS = BASE_DIR / "outputs" / "revision_figures"

TITLE = "A Cross-Flaviviral Transcriptomic Prioritization Framework for Host-Directed Therapy in Kyasanur Forest Disease"
RUNNING_TITLE = "Transcriptomic HDT Framework for KFD"


def set_margins(doc: Document) -> None:
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)


def set_base_style(doc: Document, size: int = 12) -> None:
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(size)
    style.paragraph_format.line_spacing = 2.0


def set_cell_shading(cell, color: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), color)
    tc_pr.append(shd)


def add_cited_paragraph(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph()
    parts = re.split(r"(\[\d+(?:[-,]\d+)*\])", text)
    for part in parts:
        if part.startswith("[") and part.endswith("]"):
            run = paragraph.add_run(part)
            run.font.superscript = True
        else:
            paragraph.add_run(part)


def word_count(text: str) -> int:
    return len(re.findall(r"\b\w+\b", text))


def manuscript_references() -> list[str]:
    return [
        "Work TH, Trapido H, Murthy DP, Rao RL, Bhatt PN, Kulkarni KG. Kyasanur forest disease. III. A preliminary report on the nature of the infection and clinical manifestations in man. Indian J Med Sci 1957;11:619-45.",
        "Pattnaik P. Kyasanur forest disease: an epidemiological view in India. Rev Med Virol 2006;16:151-65.",
        "Murhekar MV, Kasabi GS, Mehendale SM, Mourya DT, Yadav PD, Tandale BV, et al. On the transmission pattern of Kyasanur Forest Disease (KFD) in India. Infect Dis Poverty 2015;4:37.",
        "Holbrook MR. Kyasanur forest disease. Antiviral Res 2012;96:353-62.",
        "Kasabi GS, Murhekar MV, Sandhya VK, Raghunandan R, Kiran SK, Channabasappa GH, et al. Coverage and effectiveness of Kyasanur forest disease vaccine in Karnataka, South India, 2005-10. PLoS Negl Trop Dis 2013;7:e2025.",
        "Kaufmann SHE, Dorhoi A, Hotchkiss RS, Bartenschlager R. Host-directed therapies for bacterial and viral infections. Nat Rev Drug Discov 2018;17:35-56.",
        "Zumla A, Rao M, Wallis RS, Kaufmann SH, Rustomjee R, Mwaba P, et al. Host-directed therapies for infectious diseases: current status, recent progress, and future prospects. Lancet Infect Dis 2016;16:e47-63.",
        "Barrett T, Wilhite SE, Ledoux P, Evangelista C, Kim IF, Tomashevsky M, et al. NCBI GEO: archive for functional genomics data sets-update. Nucleic Acids Res 2013;41:D991-5.",
        "Nascimento EJ, Braga-Neto U, Calzavara-Silva CE, Gomes AL, Abath FG, Brito CA, et al. Gene expression profiling during early acute febrile stage of dengue infection can predict the disease outcome. PLoS One 2009;4:e7892.",
        "Sun P, Garcia J, Comach G, Vahey MT, Wang Z, Forshey BM, et al. Sequential waves of gene expression in patients with clinically defined dengue illnesses reveal subtle disease phases and predict disease severity. PLoS Negl Trop Dis 2013;7:e2298.",
        "Kwissa M, Nakaya HI, Onlamoon N, Wrammert J, Villinger F, Perng GC, et al. Dengue virus infection induces expansion of a CD14+CD16+ monocyte population that stimulates plasmablast differentiation. Cell Host Microbe 2014;16:115-27.",
        "Gillespie M, Jassal B, Stephan R, Milacic M, Rothfels K, Senff-Ribeiro A, et al. The Reactome pathway knowledgebase 2022. Nucleic Acids Res 2022;50:D687-92.",
        "Zdrazil B, Felix E, Hunter F, Manners EJ, Blackshaw J, Corbett S, et al. The ChEMBL Database in 2023: a drug discovery platform spanning multiple bioactivity data types and time periods. Nucleic Acids Res 2024;52:D1180-92.",
        "Simmons CP, Farrar JJ, Nguyen vV, Wills B. Dengue. N Engl J Med 2012;366:1423-32.",
        "Modhiran N, Watterson D, Muller DA, Panetta AK, Sester DP, Liu L, et al. Dengue virus NS1 protein activates cells via Toll-like receptor 4 and disrupts endothelial cell monolayer integrity. Sci Transl Med 2015;7:304ra142.",
        "Bray M. Pathogenesis of viral hemorrhagic fever. Curr Opin Immunol 2005;17:399-403.",
        "Yacoub S, Wills B. Predicting outcome from dengue. BMC Med 2014;12:147.",
        "Roberts I, Shakur H, Coats T, Hunt B, Balogun E, Barnetson L, et al. The CRASH-2 trial: a randomised controlled trial and economic evaluation of the effects of tranexamic acid on death, vascular occlusive events and transfusion requirement in bleeding trauma patients. Health Technol Assess 2013;17:1-79.",
        "Kiran SK, Padamashree S, Jayashree K, Srinath S. Health-care-seeking behaviour among Kyasanur forest disease patients in Shivamogga district, Karnataka: a cross-sectional study. Indian J Community Med 2021;46:486-9.",
    ]


def build_manuscript() -> tuple[Path, int, int]:
    targets = pd.read_csv(REV_TABLES / "kfd_revision_targets.csv")
    pathway_summary = pd.read_csv(REV_TABLES / "kfd_revision_pathway_summary.csv")
    drug_candidates = pd.read_csv(REV_TABLES / "kfd_revision_drug_candidates.csv")
    cohorts = pd.read_csv(REV_TABLES / "cohort_summary.csv")
    sensitivity = pd.read_csv(REV_TABLES / "kfd_revision_weight_sensitivity_summary.csv")

    severe_total = int(cohorts["SevereSamples"].sum())
    non_severe_total = int(cohorts["NonSevereSamples"].sum())
    total_samples = severe_total + non_severe_total

    doc = Document()
    set_margins(doc)
    set_base_style(doc)

    title = doc.add_heading("", level=0)
    run = title.add_run(TITLE)
    run.bold = True
    run.font.size = Pt(14)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    rt = doc.add_paragraph()
    rt.add_run("Running Title: ").bold = True
    rt.add_run(RUNNING_TITLE)
    doc.add_page_break()

    doc.add_heading("ABSTRACT", level=1)
    abstract_sections = [
        (
            "Background:",
            "Kyasanur Forest Disease (KFD) is a tick-borne flaviviral hemorrhagic fever endemic to the Western Ghats of India. No specific antiviral treatment exists.[1-5]",
        ),
        (
            "Objectives:",
            "To replace an opaque 'multi-omics' narrative with a reproducible transcriptomic prioritization framework for host-directed therapy (HDT) in KFD.",
        ),
        (
            "Materials and Methods:",
            f"We analyzed three public dengue severity cohorts from GEO (GSE18090, GSE43777, GSE51808; {total_samples} acute samples, {severe_total} severe and {non_severe_total} non-severe) because KFD-specific blood transcriptomes are unavailable.[8-11] A prespecified 50-gene host-response panel was rescored using an explicit formula: 0.45×omics + 0.20×tractability + 0.20×pathway relevance + 0.15×clinical-phase relevance.",
        ),
        (
            "Results:",
            "Cytokine signaling had the highest mean pathway score (0.474), followed by coagulation/fibrinolysis (0.406) and endothelial-barrier biology (0.403). The highest-ranked genes were IL1B, CXCL10, IL6, TNF, and CCL2. Endothelial and coagulation genes remained mechanistically relevant, including ANGPT2, F3, VWF, PROCR, and SERPINE1, but most showed limited cross-cohort recurrence. Rank stability remained high under alternative weighting schemes (rho 0.86-0.96). Candidate interventions were framed as hypothesis-generating, with plasma or platelet support, tranexamic acid, atorvastatin, and N-acetylcysteine shortlisted for further evaluation.",
        ),
        (
            "Conclusions:",
            "The revised study supports a transcriptomic prioritization framework, not a KFD-specific omics discovery paper. Severe flaviviral blood-response data most strongly support inflammatory signaling, while endothelial and coagulation pathways remain testable translational hypotheses for KFD that require experimental and clinical validation before use.[6,7,14-18]",
        ),
    ]
    for label, body in abstract_sections:
        paragraph = doc.add_paragraph()
        paragraph.add_run(label).bold = True
        paragraph.add_run(" " + body)

    keywords = doc.add_paragraph()
    keywords.add_run("Keywords: ").bold = True
    keywords.add_run("Kyasanur Forest Disease; host-directed therapy; transcriptomics; dengue; drug repurposing; endothelial dysfunction; coagulation")
    doc.add_page_break()

    doc.add_heading("INTRODUCTION", level=1)
    intro_paragraphs = [
        "Kyasanur Forest Disease is a tick-borne flaviviral hemorrhagic fever first recognized in Karnataka in 1957 and now reported across several Western Ghats districts.[1-4] The disease remains clinically important because it affects forest-linked rural populations, supportive care is the only established treatment, and vaccine protection is incomplete.[4,5,19]",
        "The original submission overstated the work as an 'integrated multi-omics pipeline'. That description was not justified because the study relied on transcriptomic public datasets and downstream drug-target prioritization rather than true multi-layer omics integration. The revised manuscript therefore reframes the study as a transcriptomic prioritization framework.",
        "KFD-specific blood transcriptomes are not currently available in GEO. We therefore used publicly available human dengue severity datasets as a cross-flaviviral proxy because dengue provides the best-annotated human datasets for vascular leak, thrombocytopenia, and inflammatory escalation among flaviviral illnesses.[8-11,14,15] This inference is biologically plausible but imperfect, so the revised manuscript treats all translational outputs as hypothesis-generating.",
        "The goal of the revision was not to claim KFD-specific gene discovery. Instead, we defined a reproducible 50-gene host-response panel anchored to Reactome-supported pathobiology modules relevant to hemorrhagic flaviviral disease and rescored that panel transparently against public datasets to identify HDT targets suitable for future preclinical and clinical testing.[6,7,12]",
    ]
    for paragraph in intro_paragraphs:
        add_cited_paragraph(doc, paragraph)

    doc.add_heading("MATERIALS AND METHODS", level=1)
    methods_paragraphs = [
        "Study design: This was a secondary analysis of publicly available de-identified transcriptomic datasets combined with transparent rule-based therapeutic prioritization. No human subjects were enrolled directly, and institutional ethics approval was not required for reuse of public data.",
        f"Dataset selection: Three GEO datasets containing human acute dengue samples with severe-versus-non-severe annotations were included: GSE18090 (Brazilian PBMC cohort), GSE43777 (Venezuelan longitudinal PBMC cohort; only acute-phase samples retained), and GSE51808 (Thai whole-blood cohort).[9-11] Across discovery cohorts, {severe_total} severe and {non_severe_total} non-severe acute samples were analyzed.",
        "Preprocessing and differential expression: We used the processed series-matrix files supplied by GEO submitters and analyzed each cohort separately to avoid inappropriate cross-platform normalization. Probe identifiers were mapped to official gene symbols using the corresponding GPL annotations, probe-level values were log2 transformed when the processed signal range exceeded 50, and duplicate probes were collapsed by retaining the probe with the highest mean expression. Severe-versus-non-severe contrasts were tested with Welch's t test, and Benjamini-Hochberg adjusted P values were calculated within each cohort.",
        "Gene-panel construction: The revised manuscript no longer presents the 50 genes as a KFD-specific discovery signature. Instead, the panel was prespecified from hemorrhagic-flaviviral host-response modules covering cytokine signaling, endothelial barrier biology, coagulation/fibrinolysis, platelet activation, neurological/barrier responses, and oxidative stress. Module labels were aligned to Reactome categories, and the full panel is provided in the Supplementary Materials.[12]",
        "Composite scoring: Omics score was defined as 0.65 times cross-cohort recurrence plus 0.35 times median absolute log2 fold-change. Recurrence equaled the proportion of discovery cohorts in which a gene showed same-direction differential expression at P less than or equal to 0.05 and absolute log2 fold-change greater than or equal to 0.30. Tractability score was assigned from explicit repurposing feasibility tiers (high 1.00, moderate 0.65, supportive 0.45, low 0.20). Pathway relevance and clinical-phase relevance were prespecified from KFD pathobiology. The final priority score was 0.45 times omics plus 0.20 times tractability plus 0.20 times pathway relevance plus 0.15 times phase relevance.",
        f"Weight sensitivity: To test whether the ranking depended on one arbitrary set of coefficients, we recalculated all scores under equal-weight and omics-heavy schemes. Rank correlation with the base model remained high (rho={sensitivity.iloc[0]['SpearmanRho']:.2f} for equal weighting and rho={sensitivity.iloc[1]['SpearmanRho']:.2f} for omics-heavy weighting).",
        "Drug-candidate interpretation: Candidate therapies were linked only after the target ranking was generated. The revised manuscript distinguishes supportive-care products already used in bleeding management from true repurposing hypotheses and avoids direct therapeutic recommendations for district-hospital practice in the absence of experimental or clinical validation.",
    ]
    for paragraph in methods_paragraphs:
        add_cited_paragraph(doc, paragraph)

    doc.add_heading("RESULTS", level=1)
    results_paragraphs = [
        f"Cohort composition is summarized in Figure 1. The final analysis comprised {total_samples} acute samples ({severe_total} severe, {non_severe_total} non-severe) across three geographically distinct dengue cohorts. Nominal severe-versus-non-severe differential-expression signals were detected in all three studies, but only GSE43777 retained a substantial number of within-cohort false-discovery-rate significant genes. This justified the revision strategy of using transparent cross-cohort recurrence within a prespecified panel rather than presenting an apparently de novo KFD signature.",
        "Table 1 shows the 15 highest-ranked targets. IL1B, CXCL10, and IL6 were the top three genes, followed by TNF and CCL2. Among pathway categories with immediate translational relevance to hemorrhagic disease, ANGPT2 ranked eighth, F3 tenth, VWF thirteenth, PROCR fourteenth, and SERPINE1 fifteenth. However, most endothelial and coagulation genes were prioritized because of mechanistic plausibility plus single-cohort transcriptomic support rather than robust recurrence across all datasets.",
        f"At the pathway level, cytokine signaling had the highest mean composite score ({pathway_summary.iloc[0]['MeanScore']:.3f}), followed by coagulation/fibrinolysis ({pathway_summary.iloc[1]['MeanScore']:.3f}) and endothelial-barrier biology ({pathway_summary.iloc[2]['MeanScore']:.3f}) (Table 2; Figure 4). This ranking is consistent across the text, the summary table, and the figures.",
        "The revised drug table is framed as a hypothesis-generation output rather than a recommendation list. Target-matched biologics such as anakinra and tocilizumab score highly on mechanism but are unlikely to be first-line peripheral options. Lower-cost adjunctive hypotheses include atorvastatin for endothelial stabilization, tranexamic acid for fibrinolysis-dominant bleeding phenotypes, and N-acetylcysteine for oxidative-stress modulation. Plasma and platelet support remain standard supportive care rather than discoveries of the computational pipeline.",
    ]
    for paragraph in results_paragraphs:
        add_cited_paragraph(doc, paragraph)

    doc.add_paragraph()
    cap1 = doc.add_paragraph()
    cap1.add_run("Table 1. Top 15 prioritized host-response targets.").bold = True
    table1 = doc.add_table(rows=16, cols=6)
    table1.style = "Table Grid"
    headers = ["Rank", "Gene", "Pathway", "Supporting datasets", "Composite score", "Lead hypothesis"]
    for idx, header in enumerate(headers):
        cell = table1.rows[0].cells[idx]
        cell.text = header
        cell.paragraphs[0].runs[0].bold = True
        set_cell_shading(cell, "D9E2F3")
    for row_idx, (_, row) in enumerate(targets.head(15).iterrows(), start=1):
        values = [
            str(int(row["Rank"])),
            row["GeneSymbol"],
            row["Pathway"].replace("_", " "),
            row["DatasetsSupporting"],
            f"{row['CompositeScore']:.3f}",
            row["RepurposingLead"],
        ]
        for col_idx, value in enumerate(values):
            table1.rows[row_idx].cells[col_idx].text = value

    doc.add_paragraph()
    cap2 = doc.add_paragraph()
    cap2.add_run("Table 2. Pathway-level prioritization summary.").bold = True
    table2 = doc.add_table(rows=len(pathway_summary) + 1, cols=4)
    table2.style = "Table Grid"
    headers2 = ["Pathway", "Targets", "Mean score", "SD"]
    for idx, header in enumerate(headers2):
        cell = table2.rows[0].cells[idx]
        cell.text = header
        cell.paragraphs[0].runs[0].bold = True
        set_cell_shading(cell, "D9E2F3")
    for row_idx, (_, row) in enumerate(pathway_summary.iterrows(), start=1):
        values = [
            row["Pathway"].replace("_", " "),
            str(int(row["Targets"])),
            f"{row['MeanScore']:.3f}",
            f"{row['SD']:.3f}",
        ]
        for col_idx, value in enumerate(values):
            table2.rows[row_idx].cells[col_idx].text = value

    doc.add_paragraph()
    cap3 = doc.add_paragraph()
    cap3.add_run("Table 3. Hypothesis-generating candidate interventions linked to prioritized pathways.").bold = True
    display_drugs = drug_candidates.head(8).copy()
    table3 = doc.add_table(rows=len(display_drugs) + 1, cols=5)
    table3.style = "Table Grid"
    headers3 = ["Candidate", "Target", "Pathway", "Evidence tier", "Interpretation"]
    for idx, header in enumerate(headers3):
        cell = table3.rows[0].cells[idx]
        cell.text = header
        cell.paragraphs[0].runs[0].bold = True
        set_cell_shading(cell, "D9E2F3")
    for row_idx, (_, row) in enumerate(display_drugs.iterrows(), start=1):
        values = [row["Candidate"], row["GeneSymbol"], row["Pathway"].replace("_", " "), row["EvidenceTier"], row["Rationale"]]
        for col_idx, value in enumerate(values):
            table3.rows[row_idx].cells[col_idx].text = str(value)

    doc.add_paragraph()
    figure_captions = [
        ("figure1_discovery_cohorts.png", "Figure 1. Discovery cohorts used in the revision. Bars show severe and non-severe sample counts per dataset; no inferential statistics are displayed because this panel is descriptive."),
        ("figure2_target_ranking.png", "Figure 2. Top 20 ranked targets. Horizontal bars show the composite priority score; colors denote pathway assignment."),
        ("figure3_signature_heatmap.png", "Figure 3. Cross-cohort transcriptomic evidence across the prespecified 50-gene panel. Values are cohort-specific severe-versus-non-severe log2 fold-changes."),
        ("figure4_pathway_scores.png", "Figure 4. Mean pathway scores across the full 50-gene panel. Bars show the average composite score by pathway module."),
        ("figure5_candidate_table.png", "Figure 5. Candidate interventions linked to top-ranked targets. This is a summary visualization of the hypothesis-generating shortlist rather than a clinical recommendation algorithm."),
    ]
    for filename, caption in figure_captions:
        paragraph = doc.add_paragraph()
        paragraph.add_run(caption).bold = True
        doc.add_picture(str(REV_FIGS / filename), width=Inches(5.7))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph()

    doc.add_heading("DISCUSSION", level=1)
    discussion_paragraphs = [
        "The revised analysis supports a narrower claim than the original submission. The dominant signal in the public flaviviral severity datasets was inflammatory and interferon-associated, not endothelial biology alone. This is why cytokine signaling now ranks highest at the pathway level. The revised text therefore avoids overstating endothelial pathways as the principal transcriptomic finding.",
        "Endothelial and coagulation pathways nonetheless remain translationally important for KFD. ANGPT2, F3, VWF, PROCR, and SERPINE1 remained within the top 15 targets, and those modules align with the vascular leak and bleeding phenotype that clinicians actually face in KFD.[14-18] Their evidentiary basis in this revision is weaker than that for inflammatory signaling, so they should be interpreted as clinically motivated mechanistic hypotheses rather than transcriptomically validated KFD drivers.",
        "The revised drug interpretation is intentionally conservative. Anakinra and tocilizumab emerge because they match top-ranked inflammatory targets, but they are specialist interventions and not realistic district-hospital defaults. Atorvastatin, tranexamic acid, and N-acetylcysteine remain of interest because they are lower-cost and mechanistically plausible, but the manuscript now describes them as candidates for staged preclinical or clinical evaluation rather than recommended treatment for routine use.",
        "This revision also narrows the biological extrapolation. The prior version referred broadly to viral hemorrhagic fevers and listed datasets that were not always directly relevant. The current revision limits discovery to human dengue severity cohorts, which are still imperfect proxies for KFD but are more justifiable because they share flaviviral biology and clinically relevant vascular manifestations.[9-11,14,17] Even with this narrower scope, tissue tropism, timing, and host-population differences remain major sources of uncertainty.",
    ]
    for paragraph in discussion_paragraphs:
        add_cited_paragraph(doc, paragraph)

    doc.add_heading("LIMITATIONS", level=2)
    limitation_text = (
        "No KFD-specific transcriptomic dataset was available. All omics evidence therefore comes from dengue, not KFD, and only blood-derived transcriptomes were analyzed. The 50-gene panel is prespecified and mechanistic rather than an unbiased genome-wide discovery signature. Drug-to-target links were used for prioritization only and do not establish efficacy, safety, or dosing in KFD. Experimental validation, animal models, and prospective clinical studies remain necessary before any therapeutic use."
    )
    doc.add_paragraph(limitation_text)

    doc.add_heading("CONCLUSIONS", level=1)
    conclusion_text = (
        "The revised manuscript should be interpreted as a transparent transcriptomic prioritization framework for KFD, not as a validated multi-omics discovery study. Severe flaviviral blood-response data prioritize inflammatory signaling overall, while endothelial and coagulation modules remain mechanistically and clinically relevant hypotheses for host-directed therapy development in KFD."
    )
    doc.add_paragraph(conclusion_text)

    doc.add_page_break()
    doc.add_heading("REFERENCES", level=1)
    for idx, reference in enumerate(manuscript_references(), start=1):
        paragraph = doc.add_paragraph()
        paragraph.add_run(f"{idx}. ").bold = True
        paragraph.add_run(reference)
        paragraph.paragraph_format.first_line_indent = Inches(-0.25)
        paragraph.paragraph_format.left_indent = Inches(0.25)

    output_path = MANUSCRIPT_DIR / "Manuscript_KFD_MJDYPV_Revised_Blinded.docx"
    doc.save(output_path)

    full_text = "\n".join(par.text for par in doc.paragraphs)
    total_words = word_count(full_text)
    abstract_text = " ".join(f"{label} {text}" for label, text in abstract_sections)
    abstract_words = word_count(abstract_text)
    return output_path, total_words, abstract_words


def build_title_page(total_words: int, abstract_words: int) -> Path:
    doc = Document()
    set_margins(doc)
    set_base_style(doc)

    fields = [
        ("Article Type", "Original Article"),
        ("Title", TITLE),
        ("Running Title", RUNNING_TITLE),
        ("Author", "Siddalingaiah H S"),
        ("Affiliation", "Department of Community Medicine, Shridevi Institute of Medical Sciences and Research Hospital, Tumkur, Karnataka, India"),
        ("Corresponding Author", "Dr. Siddalingaiah H S, MD; hssling@yahoo.com; +91-8941087719; ORCID 0000-0002-4771-8285"),
        ("Word Count", f"Abstract {abstract_words}; total manuscript text including references approximately {total_words}"),
        ("Tables/Figures", "3 tables and 5 figures in the blinded article; supplementary material provided separately"),
        ("Funding", "None"),
        ("Conflicts of Interest", "None declared"),
        ("Author Contributions", "Single-author study design, data analysis, interpretation, and manuscript drafting"),
        ("Ethics Statement", "Secondary analysis of publicly available de-identified datasets; ethics approval not required"),
        ("AI Disclosure", "Generative AI tools were used only for coding and editorial assistance under full author review and responsibility"),
        ("Data Availability", "All revision scripts and generated tables are contained in the submission workspace"),
    ]
    for label, value in fields:
        paragraph = doc.add_paragraph()
        paragraph.add_run(f"{label}: ").bold = True
        paragraph.add_run(value)

    output_path = MANUSCRIPT_DIR / "TitlePage_KFD_MJDYPV_Revised.docx"
    doc.save(output_path)
    return output_path


def build_cover_letter() -> Path:
    doc = Document()
    set_margins(doc)
    set_base_style(doc)

    doc.add_paragraph(datetime.now().strftime("%B %d, %Y"))
    doc.add_paragraph()
    doc.add_paragraph("The Editor-in-Chief")
    doc.add_paragraph("Medical Journal of Dr. D.Y. Patil Vidyapeeth")
    doc.add_paragraph()
    doc.add_paragraph(f"Re: revised submission of '{TITLE}'")
    doc.add_paragraph()
    body = [
        "Please find our revised original article for consideration. In response to peer review, we rebuilt the analysis rather than applying only textual edits.",
        "The revised manuscript removes the unsupported 'multi-omics' claim, replaces the opaque random-weight scoring system with an explicit deterministic framework, documents dataset preprocessing and probe-to-gene collapsing, reconciles the pathway-level narrative with the actual table values, and reframes all therapeutic outputs as hypothesis-generating.",
        "We also provide a separate response-to-reviewers table and supplementary methods/tables so that the gene-panel construction, scoring logic, and sensitivity analysis are directly auditable.",
        "This manuscript is not under consideration elsewhere. No conflicts of interest or external funding apply.",
    ]
    for paragraph in body:
        doc.add_paragraph(paragraph)
    doc.add_paragraph()
    doc.add_paragraph("Sincerely,")
    doc.add_paragraph()
    doc.add_paragraph("Dr. Siddalingaiah H S")

    output_path = MANUSCRIPT_DIR / "CoverLetter_KFD_MJDYPV_Revised.docx"
    doc.save(output_path)
    return output_path


def build_response_letter() -> Path:
    rows = [
        ("1. Methodology transparency is insufficient / gene signature derivation unclear.", "The manuscript no longer presents the panel as an unexplained KFD-specific signature. We now state explicitly that the 50-gene panel is a prespecified mechanistic host-response panel anchored to Reactome-supported modules and rescored against three public dengue severity cohorts. Preprocessing, probe collapsing, log2 transformation rules, severe-vs-nonsevere contrasts, and multiple-testing correction are described in detail.", "Methods: Dataset selection, preprocessing, gene-panel construction; Supplementary Methods and Tables S1-S2"),
        ("2. Composite scoring system is insufficiently defined.", "The previous opaque score was replaced with a deterministic formula: 0.45×omics + 0.20×tractability + 0.20×pathway relevance + 0.15×clinical-phase relevance. Each component is explicitly defined, normalized, and accompanied by a weight-sensitivity analysis.", "Methods: Composite scoring and sensitivity analysis; Supplementary Table S4"),
        ("3. Overstatement of the 'multi-omics' approach.", "Addressed by changing the title, abstract, and main text to describe a transcriptomic prioritization framework rather than an integrated multi-omics pipeline.", "Title, Abstract, Introduction, Conclusions"),
        ("4. Inconsistency between results and interpretation.", "Corrected. The revised Results section states that cytokine signaling had the highest mean pathway score, followed by coagulation/fibrinolysis and endothelial-barrier biology, exactly matching Table 2 and Figure 4.", "Results: pathway summary; Table 2; Figure 4"),
        ("5. Clinical recommendations appear premature.", "The revised manuscript removes directive treatment language. All drug outputs are now described as hypothesis-generating candidates for preclinical or clinical evaluation. Standard supportive products are clearly distinguished from computationally prioritized repurposing hypotheses.", "Abstract, Results, Discussion, Conclusions, Table 3"),
        ("6. Cross-virus extrapolation requires stronger justification.", "The revised analysis narrows the extrapolation to human dengue severity datasets only, rather than a broad mixture of viral hemorrhagic-fever datasets. The rationale and limitations of this cross-flaviviral inference are explicitly discussed.", "Introduction, Methods, Discussion, Limitations"),
        ("7. Title accuracy.", "Revised title now removes the multi-omics claim and reflects the actual study design.", "Title page and blinded manuscript title"),
        ("8. Abstract wording implied a KFD-specific signature.", "The abstract now refers to a prespecified host-response panel evaluated using public flaviviral transcriptomic data, not a KFD-specific signature.", "Abstract"),
        ("9. Reference relevance and unsupported statements.", "Irrelevant citations were removed, and the reference list was rebuilt around KFD epidemiology, dengue discovery cohorts, Reactome, ChEMBL, and host-directed therapy literature directly supporting the revised statements.", "References throughout"),
        ("10. Pathway classification requires ontology support.", "Pathway labels are now explicitly mapped to Reactome-supported modules and documented in the Methods and Supplementary Materials.", "Methods and Supplementary Table S2"),
        ("11. Vitamin K prioritization needs clarification.", "Vitamin K was removed from the high-priority shortlist. The revised candidate table instead focuses on support products, tranexamic acid, atorvastatin, and N-acetylcysteine, with appropriate cautionary framing.", "Table 3 and Discussion"),
        ("12. Figure descriptions require labels/statistical context/sample sizes.", "All revised figure captions now specify what is plotted, the sample context, and the interpretive scope of the figure.", "Figure captions 1-5"),
        ("13. Limitations should be emphasized.", "The revised paper now includes a dedicated Limitations subsection covering dataset heterogeneity, absence of KFD-specific transcriptomic data, tissue-specific uncertainty, and lack of experimental validation.", "Limitations section"),
    ]

    doc = Document()
    set_margins(doc)
    set_base_style(doc, size=11)
    heading = doc.add_heading("", level=0)
    run = heading.add_run("Response to Reviewers")
    run.bold = True
    run.font.size = Pt(14)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("Manuscript title: " + TITLE)
    doc.add_paragraph("Summary: The analysis and manuscript were rebuilt to address reproducibility, accuracy, and overstatement concerns.")

    table = doc.add_table(rows=len(rows) + 1, cols=3)
    table.style = "Table Grid"
    headers = ["Reviewer Comment", "Response", "Location in Revision"]
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        cell.text = header
        cell.paragraphs[0].runs[0].bold = True
        set_cell_shading(cell, "D9E2F3")
    for row_idx, row in enumerate(rows, start=1):
        for col_idx, value in enumerate(row):
            table.rows[row_idx].cells[col_idx].text = value

    output_path = MANUSCRIPT_DIR / "Response_to_Reviewers_KFD_MJDYPV.docx"
    doc.save(output_path)
    return output_path


def build_supplementary() -> Path:
    panel = pd.read_csv(REV_TABLES / "kfd_revision_signature.csv")
    targets = pd.read_csv(REV_TABLES / "kfd_revision_targets.csv")
    drugs = pd.read_csv(REV_TABLES / "kfd_revision_drug_candidates.csv")
    cohorts = pd.read_csv(REV_TABLES / "cohort_summary.csv")
    sensitivity = pd.read_csv(REV_TABLES / "kfd_revision_weight_sensitivity_summary.csv")

    doc = Document()
    set_margins(doc)
    set_base_style(doc, size=11)

    title = doc.add_heading("", level=0)
    run = title.add_run("Supplementary Materials")
    run.bold = True
    run.font.size = Pt(14)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(TITLE)

    doc.add_heading("Supplementary Methods", level=1)
    for item in [
        "1. GEO datasets were downloaded as processed series-matrix files from NCBI GEO. No raw CEL cross-normalization was attempted because cohorts came from different platforms and the revision focused on within-dataset severe-versus-nonsevere contrasts.",
        "2. Probe identifiers were mapped to gene symbols using the official GPL annotation tables. When multiple probes mapped to the same symbol, the probe with the highest mean expression across retained samples was selected.",
        "3. The 50-gene panel was treated as a prespecified mechanistic set rather than a de novo discovery signature. Ranking depended on transparent component scores and not on random noise terms.",
    ]:
        doc.add_paragraph(item)

    def add_df_table(df: pd.DataFrame, heading_text: str) -> None:
        doc.add_heading(heading_text, level=1)
        table = doc.add_table(rows=len(df) + 1, cols=len(df.columns))
        table.style = "Table Grid"
        for idx, column in enumerate(df.columns):
            cell = table.rows[0].cells[idx]
            cell.text = str(column)
            cell.paragraphs[0].runs[0].bold = True
            set_cell_shading(cell, "D9E2F3")
        for row_idx, (_, row) in enumerate(df.iterrows(), start=1):
            for col_idx, column in enumerate(df.columns):
                value = row[column]
                if isinstance(value, float):
                    text = f"{value:.3f}"
                else:
                    text = str(value)
                table.rows[row_idx].cells[col_idx].text = text

    add_df_table(cohorts, "Table S1. Discovery cohort summary")
    add_df_table(panel[["Gene", "GeneSymbol", "Pathway", "Phase_Relevance", "Druggability"]], "Table S2. Prespecified 50-gene host-response panel")
    add_df_table(targets[["Rank", "GeneSymbol", "Pathway", "DatasetCount", "CompositeScore", "PerDatasetLog2FC"]], "Table S3. Full ranked target panel")
    add_df_table(sensitivity, "Table S4. Weight-sensitivity analysis")
    add_df_table(drugs, "Table S5. Candidate intervention shortlist")

    output_path = MANUSCRIPT_DIR / "Supplementary_Materials_KFD_Revised.docx"
    doc.save(output_path)
    return output_path


def main() -> None:
    manuscript_path, total_words, abstract_words = build_manuscript()
    title_path = build_title_page(total_words, abstract_words)
    cover_path = build_cover_letter()
    response_path = build_response_letter()
    supplementary_path = build_supplementary()

    print("Generated revised submission package:")
    for path in [manuscript_path, title_path, cover_path, response_path, supplementary_path]:
        print(f" - {path.name}")


if __name__ == "__main__":
    main()
