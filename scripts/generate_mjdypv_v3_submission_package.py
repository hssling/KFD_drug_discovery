"""Generate the final submission package for MJDYPV."""

from __future__ import annotations

import re
from datetime import datetime
from pathlib import Path

import matplotlib.pyplot as plt
import pandas as pd
import seaborn as sns
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt


BASE_DIR = Path(__file__).resolve().parent.parent
MANUSCRIPT_DIR = BASE_DIR / "manuscripts"
REV_TABLES = BASE_DIR / "outputs" / "revision_tables"
V2_TABLES = BASE_DIR / "outputs" / "enhanced_v2_tables"
REV_FIGS = BASE_DIR / "outputs" / "revision_figures"
V2_FIGS = BASE_DIR / "outputs" / "enhanced_v2_figures"
FINAL_FIGS = BASE_DIR / "outputs" / "final_submission_figures"

TITLE = "A Cross-Flaviviral Transcriptomic Evidence and Mechanistic Prioritization Framework for Host-Directed Therapy in Kyasanur Forest Disease"
RUNNING_TITLE = "Evidence-Based HDT Framework for KFD"

FINAL_FIGS.mkdir(parents=True, exist_ok=True)


def set_margins(doc: Document) -> None:
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)


def set_base_style(doc: Document, size: int = 12, spacing: float = 2.0) -> None:
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(size)
    style.paragraph_format.line_spacing = spacing


def set_cell_shading(cell, color: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), color)
    tc_pr.append(shd)


def word_count(text: str) -> int:
    return len(re.findall(r"\b\w+\b", text))


def add_formatted_run(paragraph, text: str) -> None:
    parts = re.split(r"(\[\d+(?:[-,]\d+)*\])", text)
    for part in parts:
        if not part:
            continue
        run = paragraph.add_run(part)
        if re.fullmatch(r"\[\d+(?:[-,]\d+)*\]", part):
            run.font.superscript = True


def references() -> list[str]:
    return [
        "Work TH, Trapido H, Murthy DP, Rao RL, Bhatt PN, Kulkarni KG. Kyasanur forest disease. III. A preliminary report on the nature of the infection and clinical manifestations in man. Indian J Med Sci 1957;11:619-45.",
        "Pattnaik P. Kyasanur forest disease: an epidemiological view in India. Rev Med Virol 2006;16:151-65.",
        "Murhekar MV, Kasabi GS, Mehendale SM, Mourya DT, Yadav PD, Tandale BV, et al. On the transmission pattern of Kyasanur Forest Disease (KFD) in India. Infect Dis Poverty 2015;4:37.",
        "Holbrook MR. Kyasanur forest disease. Antiviral Res 2012;96:353-62.",
        "Kasabi GS, Murhekar MV, Sandhya VK, Raghunandan R, Kiran SK, Channabasappa GH, et al. Coverage and effectiveness of Kyasanur forest disease vaccine in Karnataka, South India, 2005-10. PLoS Negl Trop Dis 2013;7:e2025.",
        "Kaufmann SHE, Dorhoi A, Hotchkiss RS, Bartenschlager R. Host-directed therapies for bacterial and viral infections. Nat Rev Drug Discov 2018;17:35-56.",
        "Zumla A, Rao M, Wallis RS, Kaufmann SH, Rustomjee R, Mwaba P, et al. Host-directed therapies for infectious diseases: current status, recent progress, and future prospects. Lancet Infect Dis 2016;16:e47-63.",
        "Barrett T, Wilhite SE, Ledoux P, Evangelista C, Kim IF, Tomashevsky M, et al. NCBI GEO: archive for functional genomics data sets-update. Nucleic Acids Res 2013;41:D991-5.",
        "Nascimento EJ, Braga-Neto U, Calzavara-Silva CE, Gomes AL, Abath FG, Brito CA, et al. Gene expression profiling during early acute febrile stage of dengue infection can predict the disease outcome. PLoS One 2009;4:e7892.",
        "Sun P, Garcia J, Comach G, Vahey MT, Wang Z, Forshey BM, et al. Sequential waves of gene expression in patients with clinically defined dengue illnesses reveal subtle disease phases and predict disease severity. PLoS Negl Trop Dis 2013;7:e2298.",
        "Kwissa M, Nakaya HI, Onlamoon N, Wrammert J, Villinger F, Perng GC, et al. Dengue virus infection induces expansion of a CD14+CD16+ monocyte population that stimulates plasmablast differentiation. Cell Host Microbe 2014;16:115-27.",
        "Gillespie M, Jassal B, Stephan R, Milacic M, Rothfels K, Senff-Ribeiro A, et al. The Reactome pathway knowledgebase 2022. Nucleic Acids Res 2022;50:D687-92.",
        "Zdrazil B, Felix E, Hunter F, Manners EJ, Blackshaw J, Corbett S, et al. The ChEMBL Database in 2023: a drug discovery platform spanning multiple bioactivity data types and time periods. Nucleic Acids Res 2024;52:D1180-92.",
        "Simmons CP, Farrar JJ, Nguyen vV, Wills B. Dengue. N Engl J Med 2012;366:1423-32.",
        "Modhiran N, Watterson D, Muller DA, Panetta AK, Sester DP, Liu L, et al. Dengue virus NS1 protein activates cells via Toll-like receptor 4 and disrupts endothelial cell monolayer integrity. Sci Transl Med 2015;7:304ra142.",
        "Bray M. Pathogenesis of viral hemorrhagic fever. Curr Opin Immunol 2005;17:399-403.",
        "Yacoub S, Wills B. Predicting outcome from dengue. BMC Med 2014;12:147.",
        "Roberts I, Shakur H, Coats T, Hunt B, Balogun E, Barnetson L, et al. The CRASH-2 trial: a randomised controlled trial and economic evaluation of the effects of tranexamic acid on death, vascular occlusive events and transfusion requirement in bleeding trauma patients. Health Technol Assess 2013;17:1-79.",
        "Kiran SK, Padamashree S, Jayashree K, Srinath S. Health-care-seeking behaviour among Kyasanur forest disease patients in Shivamogga district, Karnataka: a cross-sectional study. Indian J Community Med 2021;46:486-9.",
    ]


def prettify_pathway(value: str) -> str:
    mapping = {
        "cytokine": "Cytokine",
        "coagulation": "Coagulation/fibrinolysis",
        "endothelial": "Endothelial barrier",
        "interferon": "Interferon",
        "neurological": "Neurological",
        "neuroprotection": "Neuroprotection",
        "oxidative": "Oxidative stress",
        "platelet": "Platelet activation",
        "cytokine_signaling": "Cytokine signaling",
        "coagulation_fibrinolysis": "Coagulation/fibrinolysis",
        "endothelial_barrier": "Endothelial barrier",
        "oxidative_stress": "Oxidative stress",
        "platelet_activation": "Platelet activation",
        "neurological_barrier": "Neurological barrier",
        "monocyte_innate_activation": "Monocyte/innate activation",
    }
    return mapping.get(value, str(value).replace("_", " ").title())


def generate_final_submission_figures() -> list[tuple[Path, str]]:
    sns.set_theme(style="whitegrid")
    cohorts = pd.read_csv(REV_TABLES / "cohort_summary.csv")
    meta = pd.read_csv(V2_TABLES / "kfd_enhanced_v2_meta_targets.csv")
    revision_targets = pd.read_csv(REV_TABLES / "kfd_revision_targets.csv")

    counts = cohorts.melt(
        id_vars=["Dataset"],
        value_vars=["SevereSamples", "NonSevereSamples"],
        var_name="Group",
        value_name="Samples",
    )
    counts["Group"] = counts["Group"].map(
        {"SevereSamples": "Severe", "NonSevereSamples": "Non-severe"}
    )
    fig, ax = plt.subplots(figsize=(8, 5))
    sns.barplot(
        data=counts,
        x="Dataset",
        y="Samples",
        hue="Group",
        ax=ax,
        palette=["#9b1d20", "#2f6690"],
    )
    ax.set_title("Figure 1. Discovery cohorts used in the analysis.")
    ax.set_xlabel("Dataset")
    ax.set_ylabel("Sample count")
    fig.tight_layout()
    fig.savefig(FINAL_FIGS / "figure1_submission.png", dpi=300)
    plt.close(fig)

    top_meta = meta.head(15).copy()
    top_meta["EvidenceTierLabel"] = top_meta["EvidenceTier"].map(
        {
            "single-cohort": "Single-cohort",
            "mechanistic-only": "Mechanistic-only",
            "cross-cohort": "Cross-cohort",
        }
    )
    fig, ax = plt.subplots(figsize=(9, 6.5))
    sns.barplot(
        data=top_meta,
        y="GeneSymbol",
        x="MetaPriority",
        hue="EvidenceTierLabel",
        dodge=False,
        palette={
            "Cross-cohort": "#1b9e77",
            "Single-cohort": "#4c78a8",
            "Mechanistic-only": "#dd8452",
        },
        ax=ax,
    )
    ax.set_title("Figure 2. Meta-priority ranking after adding pooled effects and evidence tiers.")
    ax.set_xlabel("Meta-priority score")
    ax.set_ylabel("Gene")
    ax.legend(title="Evidence tier", loc="lower right", fontsize=9)
    fig.tight_layout()
    fig.savefig(FINAL_FIGS / "figure2_submission.png", dpi=300)
    plt.close(fig)

    pathway = (
        meta.groupby("Pathway")
        .agg(
            MeanAbsEffect=("AbsRandomEffect", "mean"),
            MeanI2=("I2", "mean"),
        )
        .reset_index()
    )
    pathway["PathwayLabel"] = pathway["Pathway"].map(prettify_pathway)
    fig, ax = plt.subplots(figsize=(8.2, 5.5))
    sns.scatterplot(
        data=pathway,
        x="MeanAbsEffect",
        y="MeanI2",
        hue="PathwayLabel",
        s=140,
        ax=ax,
    )
    for _, row in pathway.iterrows():
        ax.text(
            row["MeanAbsEffect"] + 0.003,
            row["MeanI2"] + 0.4,
            row["PathwayLabel"],
            fontsize=8,
        )
    ax.set_title("Figure 3. Pathway effect size versus heterogeneity.")
    ax.set_xlabel("Mean absolute pooled effect")
    ax.set_ylabel("Mean I-squared")
    ax.legend(title="Pathway", bbox_to_anchor=(1.02, 1), loc="upper left", fontsize=8)
    fig.tight_layout()
    fig.savefig(FINAL_FIGS / "figure3_submission.png", dpi=300, bbox_inches="tight")
    plt.close(fig)

    top_revision = revision_targets.head(15).copy()
    top_revision["PathwayLabel"] = top_revision["Pathway"].map(prettify_pathway)
    fig, ax = plt.subplots(figsize=(9, 6.5))
    sns.barplot(
        data=top_revision,
        y="GeneSymbol",
        x="CompositeScore",
        hue="PathwayLabel",
        dodge=False,
        ax=ax,
    )
    ax.set_title("Figure 4. Original composite ranking retained for comparison with the meta-analytic ranking.")
    ax.set_xlabel("Composite priority score")
    ax.set_ylabel("Gene")
    ax.legend(title="Pathway", loc="lower right", fontsize=8)
    fig.tight_layout()
    fig.savefig(FINAL_FIGS / "figure4_submission.png", dpi=300)
    plt.close(fig)

    return [
        (FINAL_FIGS / "figure1_submission.png", "Figure 1. Discovery cohorts used in the analysis."),
        (FINAL_FIGS / "figure2_submission.png", "Figure 2. Meta-priority ranking after adding pooled effects and evidence tiers."),
        (FINAL_FIGS / "figure3_submission.png", "Figure 3. Pathway effect size versus heterogeneity."),
        (FINAL_FIGS / "figure4_submission.png", "Figure 4. Original composite ranking retained for comparison with the meta-analytic ranking."),
    ]


def build_blinded_manuscript() -> tuple[Path, int, int]:
    meta = pd.read_csv(V2_TABLES / "kfd_enhanced_v2_meta_targets.csv")
    transl = pd.read_csv(V2_TABLES / "kfd_enhanced_v2_translational_targets.csv")
    cohorts = pd.read_csv(REV_TABLES / "cohort_summary.csv")
    revision_drugs = pd.read_csv(REV_TABLES / "kfd_revision_drug_candidates.csv")
    single = int((meta["EvidenceTier"] == "single-cohort").sum())
    mech = int((meta["EvidenceTier"] == "mechanistic-only").sum())
    total = int(cohorts["SevereSamples"].sum() + cohorts["NonSevereSamples"].sum())
    severe = int(cohorts["SevereSamples"].sum())
    non_severe = int(cohorts["NonSevereSamples"].sum())

    doc = Document()
    set_margins(doc)
    set_base_style(doc)

    title = doc.add_heading("", level=0)
    run = title.add_run(TITLE)
    run.bold = True
    run.font.size = Pt(14)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = doc.add_paragraph()
    p.add_run("Running Title: ").bold = True
    p.add_run(RUNNING_TITLE)
    doc.add_page_break()

    doc.add_heading("ABSTRACT", level=1)
    abstract_sections = [
        ("Background:", "Kyasanur Forest Disease lacks disease-specific transcriptomic datasets and specific therapy, forcing host-directed therapeutic work to rely on proxy flaviviral data."),
        ("Objectives:", "To strengthen a transcriptomic prioritization framework for KFD by adding pooled effect estimates, confidence intervals, heterogeneity metrics, and stricter evidence grading."),
        ("Materials and Methods:", f"We reanalyzed three public dengue severity cohorts from GEO (GSE18090, GSE43777, GSE51808; {total} acute samples, {severe} severe and {non_severe} non-severe). A prespecified 50-gene host-response panel was scored using a deterministic framework and then supplemented with random-effects meta-analysis, 95% confidence intervals, and evidence-tier classification."),
        ("Results:", f"The strongest transcriptomic evidence remained inflammatory. No gene reached a strict cross-cohort evidence tier; {single} genes had single-cohort nominal support and {mech} remained mechanistic-only. The highest meta-priority genes were CXCL10, IL6, STAT1, IFNB1, and TEK. Endothelial/coagulation genes such as ANGPT2, F3, VWF, and SERPINE1 remained biologically relevant but weakly supported transcriptomically."),
        ("Conclusions:", "This framing separates transcriptomic evidence from mechanistic plausibility. Inflammatory targets are the best-supported findings in current public data, while endothelial and coagulation pathways remain clinically relevant KFD hypotheses that require KFD-specific validation before therapeutic inference."),
    ]
    for label, text in abstract_sections:
        p = doc.add_paragraph()
        p.add_run(label).bold = True
        add_formatted_run(p, " " + text)
    p = doc.add_paragraph()
    p.add_run("Keywords: ").bold = True
    p.add_run("Kyasanur Forest Disease; host-directed therapy; transcriptomics; meta-analysis; flavivirus; endothelial dysfunction; coagulation")
    doc.add_page_break()

    doc.add_heading("INTRODUCTION", level=1)
    for text in [
        "Kyasanur Forest Disease (KFD) is a tick-borne flaviviral hemorrhagic fever first recognized in Karnataka and remains an important clinical and public health problem in affected districts of southern India.[1-4] Patients can present with intense febrile illness, thrombocytopenia, bleeding manifestations, and in some cases neurological complications, yet no disease-specific antiviral treatment is available and vaccine protection remains incomplete in field conditions.[4,5]",
        "This therapeutic gap makes host-directed therapy an attractive idea for KFD because the clinically important manifestations of severe disease are likely driven by host inflammatory, vascular, and hemostatic pathways rather than by viral replication alone.[6,7] At the same time, host-directed therapy studies are easy to overstate if biological plausibility is treated as equivalent to reproducible molecular evidence.",
        "A major limitation in KFD research is the absence of publicly available KFD-specific blood transcriptomic datasets. We therefore used severe-versus-non-severe human dengue cohorts as a cross-flaviviral proxy because they provide the most accessible human transcriptomic data linked to severe disease, vascular leak, and hemorrhagic manifestations among related flaviviral infections.[8-11] This strategy is biologically reasonable for hypothesis generation, but it requires explicit caution because host responses vary across viruses, tissues, and stages of disease.",
        "The aim of this revised manuscript was not to claim a KFD-specific discovery signature. Instead, we evaluated a prespecified 50-gene host-response panel mapped to biologically relevant modules using Reactome-informed pathway categories and then added a random-effects meta-analytic layer to quantify pooled effect size, uncertainty, and heterogeneity.[12] The purpose was to distinguish transcriptomically supported priorities from clinically important mechanistic hypotheses and to present the resulting intervention shortlist with appropriate restraint.",
    ]:
        p = doc.add_paragraph()
        add_formatted_run(p, text)

    doc.add_heading("MATERIALS AND METHODS", level=1)
    doc.add_heading("Data sources and cohort selection", level=2)
    for text in [
        f"Three public GEO datasets with acute human dengue samples and severe-versus-non-severe annotations were included: GSE18090, GSE43777, and GSE51808.[8-11] Together they contributed {total} acute samples, including {severe} severe and {non_severe} non-severe samples. Cohort composition is shown in Figure 1.",
        "Only acute-phase samples were analyzed. We intentionally avoided pooled cross-platform normalization because the included cohorts differed in platform and specimen type. Instead, each dataset was processed within cohort, and only gene-level summary statistics were carried forward into the cross-cohort framework.",
    ]:
        p = doc.add_paragraph()
        add_formatted_run(p, text)

    doc.add_heading("Preprocessing and within-cohort analysis", level=2)
    for text in [
        "Series-matrix files and platform annotations were obtained from GEO.[8] Probe identifiers were mapped to gene symbols using platform annotation files, duplicate probes were collapsed conservatively to a single representative probe per gene, and expression values were log-transformed when required. This workflow was chosen to maximize transparency and reproducibility while minimizing additional modelling assumptions.",
        "Within each cohort, severe and non-severe samples were compared at the gene level using two-sample statistical testing. For every prespecified gene, we retained the cohort-specific log2 fold-change, nominal P value, and false-discovery-rate adjusted P value. These cohort-specific results formed the basis of both the deterministic prioritization layer and the added meta-analysis.",
    ]:
        p = doc.add_paragraph()
        add_formatted_run(p, text)

    doc.add_heading("Prespecified host-response panel and deterministic prioritization", level=2)
    for text in [
        "The 50-gene panel was not treated as a de novo KFD discovery signature. Instead, it was defined as a mechanistic host-response panel spanning cytokine signaling, interferon biology, endothelial barrier regulation, coagulation and fibrinolysis, platelet activation, oxidative stress, and neurological or barrier-related pathways. Pathway grouping was anchored to Reactome-supported biological modules to improve biological consistency.[12]",
        "We retained the deterministic prioritization framework from the earlier revision so that mechanistic and translational context would not be lost when the statistical layer was tightened. The deterministic score combined transcriptomic support, pathway relevance, disease-phase relevance, and tractability of host-directed modulation. Tractability was informed by the presence of plausible repurposing leads and by drug-discovery resources such as ChEMBL.[13]",
    ]:
        p = doc.add_paragraph()
        add_formatted_run(p, text)

    doc.add_heading("Meta-analysis and evidence-tier assignment", level=2)
    for text in [
        "To strengthen statistical rigor, each panel gene was additionally evaluated using random-effects meta-analysis based on cohort-level effect sizes and approximated standard errors derived from log2 fold-changes and two-sided P values. For each gene we report pooled effect size, 95% confidence interval, pooled P value, and heterogeneity measured as I-squared.",
        "Genes were classified into three evidence tiers. Cross-cohort support required more than one nominally supporting cohort together with a significant pooled effect. Single-cohort support required one nominally supporting cohort. Mechanistic-only denotes genes retained because of biological relevance but lacking recurrent nominal transcriptomic support. This framework was designed to separate molecular evidence from pathway plausibility rather than merge them into a single unsupported claim.",
    ]:
        p = doc.add_paragraph()
        add_formatted_run(p, text)

    doc.add_heading("RESULTS", level=1)
    doc.add_heading("Cohort composition and analytic context", level=2)
    for text in [
        f"The three cohorts were heterogeneous in sample source, geography, and severe-case representation, but all contributed severe-versus-non-severe information. GSE18090 included 10 severe and 8 non-severe PBMC samples, GSE51808 included 10 severe and 18 non-severe whole-blood samples, and GSE43777 included 37 severe and 39 non-severe PBMC samples. Figure 1 shows that GSE43777 contributed the largest share of analyzed samples.",
        "Nominal differential-expression burden varied substantially across cohorts, indicating that the proxy data are informative but not uniform. This variability supports the decision to carry forward gene-level summary statistics rather than directly pool expression matrices across studies.",
    ]:
        p = doc.add_paragraph()
        add_formatted_run(p, text)

    doc.add_heading("Evidence-tier distribution and top-ranked genes", level=2)
    for text in [
        f"No gene reached the strict cross-cohort tier. {single} genes met the single-cohort tier and {mech} remained mechanistic-only after pooled evaluation. This is the most important empirical result of the revised study: current public proxy data can support cautious prioritization, but they do not justify claims of a stable cross-flaviviral KFD-like consensus signature.",
        "The highest-ranked genes by the combined meta-priority framework were CXCL10, IL6, STAT1, IFNB1, and TEK (Table 1; Figure 2). These genes cluster mainly within inflammatory and interferon-related biology. However, the confidence intervals for several top-ranked genes remained wide and heterogeneity was often moderate to high, which is why the manuscript treats them as evidence-supported priorities rather than validated biomarkers.",
        "Table 1 therefore serves two purposes. It identifies the strongest current targets under the final framework, and it makes visible the uncertainty attached to each rank. The table should not be interpreted as a fixed disease signature but as a transparent ranking of candidates under limited-data conditions.",
    ]:
        p = doc.add_paragraph()
        add_formatted_run(p, text)

    doc.add_heading("Pathway-level findings", level=2)
    for text in [
        "At the pathway level, cytokine signaling had the highest mean deterministic score, followed by coagulation/fibrinolysis and endothelial barrier biology. Figure 3 places those pathway-level tendencies in the context of pooled effect size and heterogeneity, and Figure 4 shows the original composite ranking retained for comparison with the final meta-priority hierarchy.",
        "This pathway-level view resolves a key interpretive issue from earlier drafts. Inflammatory pathways carry the strongest transcriptomic signal in the available public data, whereas endothelial and hemostatic pathways remain prominent mainly because they are central to the known clinical pathophysiology of hemorrhagic disease and retain some degree of supporting signal, even if that signal is not sufficiently recurrent to support strong transcriptomic claims.",
    ]:
        p = doc.add_paragraph()
        add_formatted_run(p, text)

    doc.add_heading("Translationally relevant targets and intervention shortlist", level=2)
    for text in [
        "The translational subset is summarized in Table 2. IL6, TNF, IL1B, BDNF, and SERPINE1 retained single-cohort support under the final evidence framework, whereas ANGPT2, VWF, and F3 remained biologically important but fell into the mechanistic-only category. This pattern reinforces the central conclusion that mechanistic importance and transcriptomic recurrence are not interchangeable.",
        "Table 3 presents the intervention shortlist as a hypothesis-generating output rather than a recommendation for use. The shortlist intentionally includes both supportive-care aligned strategies and exploratory host-directed candidates mapped to prioritized pathways. Interventions linked to inflammatory targets are currently better grounded in transcriptomic evidence, whereas endothelial and coagulation-directed options are retained because of disease plausibility and translational relevance rather than because the molecular evidence is already strong.",
    ]:
        p = doc.add_paragraph()
        add_formatted_run(p, text)

    top_meta = meta.head(12).copy()
    p = doc.add_paragraph()
    p.add_run("Table 1. Top 12 genes ranked by meta-priority.").bold = True
    t1 = doc.add_table(rows=len(top_meta) + 1, cols=7)
    t1.style = "Table Grid"
    h1 = ["Meta rank", "Gene", "Evidence tier", "Support count", "Pooled effect", "95% CI", "I2"]
    for i, h in enumerate(h1):
        cell = t1.rows[0].cells[i]
        cell.text = h
        cell.paragraphs[0].runs[0].bold = True
        set_cell_shading(cell, "D9E2F3")
    for i, (_, row) in enumerate(top_meta.iterrows(), start=1):
        vals = [
            int(row["MetaRank"]), row["GeneSymbol"], row["EvidenceTier"], int(row["NominalSupportCount"]),
            f"{row['RandomEffect']:.2f}", f"{row['Lower95CI']:.2f} to {row['Upper95CI']:.2f}", f"{row['I2']:.1f}"
        ]
        for j, v in enumerate(vals):
            t1.rows[i].cells[j].text = str(v)

    p = doc.add_paragraph()
    p.add_run("Table 2. Translationally relevant targets with evidence tier and interpretation.").bold = True
    t2 = doc.add_table(rows=len(transl) + 1, cols=7)
    t2.style = "Table Grid"
    h2 = ["Meta rank", "Gene", "Pathway", "Evidence tier", "Support", "Pooled effect", "Interpretation"]
    for i, h in enumerate(h2):
        cell = t2.rows[0].cells[i]
        cell.text = h
        cell.paragraphs[0].runs[0].bold = True
        set_cell_shading(cell, "D9E2F3")
    interp = {
        "single-cohort": "Transcriptomic signal present but not recurrent",
        "mechanistic-only": "Mechanistically plausible but weak transcriptomic support",
        "cross-cohort": "Recurrent transcriptomic support",
    }
    for i, (_, row) in enumerate(transl.iterrows(), start=1):
        vals = [
            int(row["MetaRank"]), row["GeneSymbol"], row["Pathway"], row["EvidenceTier"],
            int(row["NominalSupportCount"]), f"{row['RandomEffect']:.2f}", interp[row["EvidenceTier"]]
        ]
        for j, v in enumerate(vals):
            t2.rows[i].cells[j].text = str(v)

    shortlist = revision_drugs[revision_drugs["GeneSymbol"].isin(["IL1B", "IL6", "TNF", "ANGPT2", "SERPINE1", "HMOX1", "F3", "VWF"])].copy()
    p = doc.add_paragraph()
    p.add_run("Table 3. Hypothesis-generating intervention shortlist under the final evidence framework.").bold = True
    t3 = doc.add_table(rows=len(shortlist) + 1, cols=5)
    t3.style = "Table Grid"
    h3 = ["Candidate", "Target", "Pathway", "Evidence tier", "Use in manuscript"]
    tier_lookup = meta.set_index("GeneSymbol")["EvidenceTier"].to_dict()
    for i, h in enumerate(h3):
        cell = t3.rows[0].cells[i]
        cell.text = h
        cell.paragraphs[0].runs[0].bold = True
        set_cell_shading(cell, "D9E2F3")
    for i, (_, row) in enumerate(shortlist.iterrows(), start=1):
        vals = [row["Candidate"], row["GeneSymbol"], row["Pathway"], tier_lookup.get(row["GeneSymbol"], "n/a"), "Hypothesis-generating only"]
        for j, v in enumerate(vals):
            t3.rows[i].cells[j].text = str(v)

    figures = generate_final_submission_figures()
    for path, caption in figures:
        p = doc.add_paragraph()
        p.add_run(caption).bold = True
        doc.add_picture(str(path), width=Inches(5.7))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_heading("DISCUSSION", level=1)
    for text in [
        "The principal strength of the revised manuscript is that it separates transcriptomic support from mechanistic plausibility instead of blending them into a single overconfident ranking. Once pooled effects, confidence intervals, and heterogeneity are considered together, inflammatory and interferon-related genes emerge as the most defensible priorities in the available proxy data.",
        "That does not make endothelial and coagulation biology unimportant. On the contrary, these pathways remain highly relevant to how severe KFD is understood clinically. What changes in the revised manuscript is the level of certainty attached to them. ANGPT2, VWF, F3, and related genes are retained because they are mechanistically credible in hemorrhagic disease, but they are no longer described as though they were recurrent transcriptomic drivers in the current public datasets.[14-17]",
        "This distinction matters for host-directed therapy research more broadly.[6,7] A pathway can be biologically attractive and clinically actionable in principle, yet still lack reproducible molecular support in currently available data. By making that boundary explicit, the manuscript becomes more reliable and more useful to readers who want to understand what is supported now and what still requires validation.",
        "The intervention shortlist should therefore be read as a staged translational agenda. Supportive-care elements remain closest to current practice, while repurposing candidates such as atorvastatin, tranexamic acid, and N-acetylcysteine should be considered priorities for preclinical testing, biomarker-linked observational work, or carefully designed early clinical evaluation rather than routine use. For tranexamic acid in particular, the present rationale should be interpreted in the context of broader hemorrhage literature and not as direct efficacy evidence for KFD.[18]",
        "The rural and forest-linked setting of KFD also remains relevant when considering translational usefulness. Any future intervention strategy has to be realistic for district-level care, delayed presentation, and healthcare-seeking patterns in endemic areas.[19] This practical dimension explains why low-cost interventions may still be worth discussing even when their supporting molecular evidence is weaker than that of inflammatory targets.",
    ]:
        p = doc.add_paragraph()
        add_formatted_run(p, text)

    doc.add_heading("LIMITATIONS", level=2)
    p = doc.add_paragraph()
    add_formatted_run(
        p,
        "No KFD-specific transcriptomic data were available. The current evidence therefore depends on dengue proxy cohorts, blood-derived transcriptomes only, and approximate variance reconstruction for the random-effects meta-analysis. The data support prioritization under uncertainty, not therapeutic validation.",
    )

    doc.add_heading("CONCLUSIONS", level=1)
    p = doc.add_paragraph()
    add_formatted_run(
        p,
        "A transparent transcriptomic prioritization framework can still be useful for KFD despite the absence of disease-specific transcriptomic datasets, but only if the outputs are interpreted with strict caution. In the current public proxy data, inflammatory targets have the strongest molecular support, whereas endothelial and coagulation pathways remain clinically meaningful mechanistic hypotheses that require KFD-specific validation before therapeutic inference.",
    )

    doc.add_page_break()
    doc.add_heading("REFERENCES", level=1)
    for idx, ref in enumerate(references(), start=1):
        p = doc.add_paragraph()
        p.add_run(f"{idx}. ").bold = True
        p.add_run(ref)
        p.paragraph_format.first_line_indent = Inches(-0.25)
        p.paragraph_format.left_indent = Inches(0.25)

    out = MANUSCRIPT_DIR / "Manuscript_KFD_MJDYPV_Final_Blinded.docx"
    doc.save(out)
    text = "\n".join(p.text for p in doc.paragraphs)
    abstract_text = " ".join(f"{a} {b}" for a, b in abstract_sections)
    return out, word_count(text), word_count(abstract_text)


def build_title_page(total_words: int, abstract_words: int) -> Path:
    doc = Document()
    set_margins(doc)
    set_base_style(doc)
    fields = [
        ("Article Type", "Original Article"),
        ("Title", TITLE),
        ("Running Title", RUNNING_TITLE),
        ("Author", "Siddalingaiah H S"),
        ("Affiliation", "Department of Community Medicine, Shridevi Institute of Medical Sciences and Research Hospital, Tumkur, Karnataka, India"),
        ("Corresponding Author", "Dr. Siddalingaiah H S, MD; hssling@yahoo.com; +91-8941087719; ORCID 0000-0002-4771-8285"),
        ("Word Count", f"Abstract {abstract_words}; total manuscript text including references approximately {total_words}"),
        ("Tables/Figures", "3 tables and 4 figures in the blinded article; supplementary material provided separately"),
        ("Funding", "None"),
        ("Conflicts of Interest", "None declared"),
        ("Ethics Statement", "Secondary analysis of publicly available de-identified datasets; ethics approval not required"),
        ("Revision Note", "Final submission package integrates transcriptomic prioritization with meta-analysis and uncertainty grading"),
    ]
    for label, value in fields:
        p = doc.add_paragraph()
        p.add_run(f"{label}: ").bold = True
        p.add_run(value)
    out = MANUSCRIPT_DIR / "TitlePage_KFD_MJDYPV_Final.docx"
    doc.save(out)
    return out


def build_cover_letter() -> Path:
    doc = Document()
    set_margins(doc)
    set_base_style(doc)
    doc.add_paragraph(datetime.now().strftime("%B %d, %Y"))
    doc.add_paragraph()
    doc.add_paragraph("The Editor-in-Chief")
    doc.add_paragraph("Medical Journal of Dr. D.Y. Patil Vidyapeeth")
    doc.add_paragraph()
    body = [
        f"We submit the enclosed manuscript entitled '{TITLE}' as an original article.",
        "The revised manuscript retains the deterministic transcriptomic prioritization framework and adds a stricter statistical layer, including random-effects meta-analysis, confidence intervals, heterogeneity metrics, and explicit evidence-tier grading.",
        "The key value of the current submission is its improved scientific restraint. The manuscript now distinguishes transcriptomically supported inflammatory findings from endothelial and coagulation hypotheses that remain clinically relevant but mechanistically prioritized under uncertainty.",
        "A point-by-point response table and supplementary materials are included. The manuscript is not under consideration elsewhere. No conflicts of interest or external funding apply.",
    ]
    for para in body:
        doc.add_paragraph(para)
    doc.add_paragraph()
    doc.add_paragraph("Sincerely,")
    doc.add_paragraph()
    doc.add_paragraph("Dr. Siddalingaiah H S")
    out = MANUSCRIPT_DIR / "CoverLetter_KFD_MJDYPV_Final.docx"
    doc.save(out)
    return out


def build_response_letter() -> Path:
    rows = [
        ("1. Methodology transparency was insufficient.", "We rebuilt the computational workflow to make each analytical step explicit. The manuscript now describes cohort selection, preprocessing, probe-to-gene collapsing, within-cohort differential-expression analysis, the prespecified 50-gene host-response panel, the deterministic prioritization formula, and the added random-effects meta-analysis with confidence intervals and heterogeneity metrics.", "Methods; Supplementary Tables S1-S6"),
        ("2. Gene signature derivation appeared arbitrary.", "We clarified that the 50 genes are not presented as a de novo KFD-specific discovery signature. Instead, they constitute a prespecified mechanistic host-response panel spanning cytokine, endothelial, coagulation, platelet, neurological, and oxidative-stress biology. The manuscript now separates panel construction from transcriptomic evidence strength.", "Abstract, Methods, Discussion, Supplementary Table S2"),
        ("3. Composite score needed clearer definition.", "We retained the explicit deterministic prioritization framework and further strengthened it by adding an orthogonal random-effects meta-analysis layer. This allows readers to distinguish mechanistic prioritization from pooled transcriptomic evidence and uncertainty.", "Methods; Results; Supplementary Tables S2-S6"),
        ("4. The original multi-omics framing was overstated.", "We removed the unsupported multi-omics framing throughout and now describe the study as a cross-flaviviral transcriptomic evidence and mechanistic prioritization framework.", "Title, Abstract, Introduction, Conclusions"),
        ("5. Results and interpretation were inconsistent.", "We reconciled the narrative with the underlying data. The final manuscript states that inflammatory targets have the strongest transcriptomic support, while endothelial and coagulation targets remain clinically meaningful but mechanistic hypotheses with weaker transcriptomic backing in the current datasets.", "Results, Discussion, Conclusions, Tables 1-3"),
        ("6. Clinical recommendations were premature.", "We removed directive therapeutic language. The intervention shortlist is now explicitly labeled hypothesis-generating only, and supportive-care products are distinguished from exploratory repurposing candidates.", "Abstract, Table 3, Discussion, Conclusions"),
        ("7. Cross-virus extrapolation required stronger justification.", "We narrowed the public discovery base to human dengue severity cohorts and added an explicit uncertainty framework. The manuscript now states clearly that current proxy datasets are insufficient for strong KFD-specific molecular claims and that disease-specific validation remains necessary.", "Introduction, Methods, Limitations, Conclusions"),
        ("8. Pathway classification required support.", "We retained pathway mapping grounded in established host-response biology and documented the categorized panel transparently in the manuscript and supplementary materials.", "Methods; Supplementary Table S2"),
        ("9. Tables and figures needed closer agreement with claims.", "All tables and figures were regenerated directly from verified output files and re-audited against the manuscript text. Figure graphics now use publication-facing numbering and titles without internal draft labels, and the supplementary tables were reformatted into compact summaries plus separate evidence-detail tables for readability.", "Tables 1-3; Figures 1-4; Supplementary Tables S1-S6"),
        ("10. Limitations required more emphasis.", "We expanded the limitations substantially. The manuscript now states that no gene reached strict cross-cohort support under the current public data base, that several vascular/coagulation targets are mechanistic rather than recurrent transcriptomic findings, and that KFD-specific data are still required for validation.", "Abstract, Results, Limitations, Conclusions"),
    ]
    doc = Document()
    set_margins(doc)
    set_base_style(doc, size=11, spacing=1.5)
    h = doc.add_heading("", level=0)
    r = h.add_run("Response to Reviewers")
    r.bold = True
    r.font.size = Pt(14)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("Manuscript title: " + TITLE)
    table = doc.add_table(rows=len(rows) + 1, cols=3)
    table.style = "Table Grid"
    headers = ["Reviewer Concern", "Response", "Location in Revised Manuscript"]
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        cell.paragraphs[0].runs[0].bold = True
        set_cell_shading(cell, "D9E2F3")
    for i, row in enumerate(rows, start=1):
        for j, value in enumerate(row):
            table.rows[i].cells[j].text = value
    out = MANUSCRIPT_DIR / "Response_to_Reviewers_KFD_MJDYPV_Final.docx"
    doc.save(out)
    return out


def build_supplementary() -> Path:
    meta = pd.read_csv(V2_TABLES / "kfd_enhanced_v2_meta_targets.csv")
    transl = pd.read_csv(V2_TABLES / "kfd_enhanced_v2_translational_targets.csv")
    evidence = pd.read_csv(V2_TABLES / "kfd_enhanced_v2_evidence_summary.csv")
    cohorts = pd.read_csv(REV_TABLES / "cohort_summary.csv")
    revision_targets = pd.read_csv(REV_TABLES / "kfd_revision_targets.csv")

    doc = Document()
    set_margins(doc)
    set_base_style(doc, size=11, spacing=1.3)
    h = doc.add_heading("", level=0)
    r = h.add_run("Supplementary Materials")
    r.bold = True
    r.font.size = Pt(14)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(TITLE)

    def add_df(title: str, df: pd.DataFrame) -> None:
        doc.add_heading(title, level=1)
        table = doc.add_table(rows=len(df) + 1, cols=len(df.columns))
        table.style = "Table Grid"
        for i, col in enumerate(df.columns):
            cell = table.rows[0].cells[i]
            cell.text = str(col)
            cell.paragraphs[0].runs[0].bold = True
            set_cell_shading(cell, "D9E2F3")
        for i, (_, row) in enumerate(df.iterrows(), start=1):
            for j, col in enumerate(df.columns):
                val = row[col]
                if isinstance(val, float):
                    txt = f"{val:.3f}"
                else:
                    txt = str(val)
                table.rows[i].cells[j].text = txt

    meta_summary = meta[
        [
            "MetaRank",
            "GeneSymbol",
            "Pathway",
            "EvidenceTier",
            "NominalSupportCount",
            "RandomEffect",
            "Lower95CI",
            "Upper95CI",
            "PooledPValue",
            "I2",
        ]
    ].copy()
    meta_summary["Pathway"] = meta_summary["Pathway"].map(prettify_pathway)
    meta_summary["Pooled effect (95% CI)"] = meta_summary.apply(
        lambda row: f"{row['RandomEffect']:.2f} ({row['Lower95CI']:.2f} to {row['Upper95CI']:.2f})",
        axis=1,
    )
    meta_summary["Pooled P value"] = meta_summary["PooledPValue"].map(lambda x: f"{x:.3g}")
    meta_summary["Mean I-squared"] = meta_summary["I2"].map(lambda x: f"{x:.1f}")
    meta_summary = meta_summary[
        [
            "MetaRank",
            "GeneSymbol",
            "Pathway",
            "EvidenceTier",
            "NominalSupportCount",
            "Pooled effect (95% CI)",
            "Pooled P value",
            "Mean I-squared",
        ]
    ]

    meta_detail = meta[["MetaRank", "GeneSymbol", "PerStudyEffects"]].copy()
    meta_detail["PerStudyEffects"] = meta_detail["PerStudyEffects"].str.replace("; ", "\n", regex=False)

    transl_summary = transl.copy()
    transl_summary["Pathway"] = transl_summary["Pathway"].map(prettify_pathway)
    transl_summary["Pooled effect (95% CI)"] = transl_summary.apply(
        lambda row: f"{row['RandomEffect']:.2f} ({row['Lower95CI']:.2f} to {row['Upper95CI']:.2f})",
        axis=1,
    )
    transl_summary = transl_summary[
        [
            "MetaRank",
            "GeneSymbol",
            "Pathway",
            "EvidenceTier",
            "NominalSupportCount",
            "Pooled effect (95% CI)",
            "I2",
        ]
    ].copy()
    transl_summary["I2"] = transl_summary["I2"].map(lambda x: f"{x:.1f}")

    transl_detail = transl[["MetaRank", "GeneSymbol", "PerStudyEffects"]].copy()
    transl_detail["PerStudyEffects"] = transl_detail["PerStudyEffects"].str.replace("; ", "\n", regex=False)

    doc.add_heading("Supplementary Methods", level=1)
    for text in [
        "Processed GEO series-matrix files were analyzed within cohort to avoid inappropriate cross-platform normalization.",
        "The 50-gene panel was retained as a prespecified mechanistic set.",
        "Random-effects meta-analysis was added to quantify pooled effect size and uncertainty for each panel gene.",
    ]:
        doc.add_paragraph(text)

    add_df("Table S1. Discovery cohort summary", cohorts)
    add_df("Table S2. Full meta-analysis target summary", meta_summary)
    add_df("Table S3. Cohort-specific effect details for the 50-gene panel", meta_detail)
    add_df("Table S4. Translational target summary", transl_summary)
    add_df("Table S5. Cohort-specific evidence details for translational targets", transl_detail)
    add_df("Table S6. Evidence-tier summary by pathway", evidence)
    add_df("Table S7. Original revision ranking retained for comparison", revision_targets[["Rank", "GeneSymbol", "Pathway", "CompositeScore", "DatasetsSupporting"]])

    out = MANUSCRIPT_DIR / "Supplementary_Materials_KFD_Final.docx"
    doc.save(out)
    return out


def main() -> None:
    generate_final_submission_figures()
    manuscript, total_words, abstract_words = build_blinded_manuscript()
    title = build_title_page(total_words, abstract_words)
    cover = build_cover_letter()
    response = build_response_letter()
    supp = build_supplementary()
    print("Generated final submission package:")
    for file in [manuscript, title, cover, response, supp]:
        print(f" - {file.name}")


if __name__ == "__main__":
    main()
