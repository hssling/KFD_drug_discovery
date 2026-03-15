"""Generate alternate layout assets: tables after references and separate figures DOCX."""

from __future__ import annotations

import re
from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt


BASE_DIR = Path(__file__).resolve().parent.parent
MANUSCRIPT_DIR = BASE_DIR / "manuscripts"
REV_TABLES = BASE_DIR / "outputs" / "revision_tables"
V2_TABLES = BASE_DIR / "outputs" / "enhanced_v2_tables"
REV_FIGS = BASE_DIR / "outputs" / "revision_figures"
V2_FIGS = BASE_DIR / "outputs" / "enhanced_v2_figures"
FINAL_FIGS = BASE_DIR / "outputs" / "final_submission_figures"


def set_margins(doc: Document) -> None:
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)


def set_style(doc: Document, size: int = 12) -> None:
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(size)
    style.paragraph_format.line_spacing = 2.0


def set_cell_shading(cell, color: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), color)
    tc_pr.append(shd)


def add_formatted_run(paragraph, text: str) -> None:
    parts = re.split(r"(\[\d+(?:[-,]\d+)*\])", text)
    for part in parts:
        if not part:
            continue
        run = paragraph.add_run(part)
        if re.fullmatch(r"\[\d+(?:[-,]\d+)*\]", part):
            run.font.superscript = True


def clone_paragraph(dest_doc: Document, source_para) -> None:
    new_para = dest_doc.add_paragraph()
    if source_para.style is not None:
        new_para.style = source_para.style
    new_para.alignment = source_para.alignment
    new_para.paragraph_format.left_indent = source_para.paragraph_format.left_indent
    new_para.paragraph_format.right_indent = source_para.paragraph_format.right_indent
    new_para.paragraph_format.first_line_indent = source_para.paragraph_format.first_line_indent
    new_para.paragraph_format.space_before = source_para.paragraph_format.space_before
    new_para.paragraph_format.space_after = source_para.paragraph_format.space_after
    new_para.paragraph_format.line_spacing = source_para.paragraph_format.line_spacing
    for run in source_para.runs:
        new_run = new_para.add_run(run.text)
        new_run.bold = run.bold
        new_run.italic = run.italic
        new_run.underline = run.underline
        new_run.font.superscript = run.font.superscript
        if run.font.name:
            new_run.font.name = run.font.name
        if run.font.size:
            new_run.font.size = run.font.size


def build_tables_after_references_variant() -> Path:
    source = Document(MANUSCRIPT_DIR / "Manuscript_KFD_MJDYPV_Final_Blinded.docx")
    meta = pd.read_csv(V2_TABLES / "kfd_enhanced_v2_meta_targets.csv")
    transl = pd.read_csv(V2_TABLES / "kfd_enhanced_v2_translational_targets.csv")
    revision_drugs = pd.read_csv(REV_TABLES / "kfd_revision_drug_candidates.csv")
    tier_lookup = meta.set_index("GeneSymbol")["EvidenceTier"].to_dict()

    doc = Document()
    set_margins(doc)
    set_style(doc)

    for para in source.paragraphs:
        # Skip inline figure-only paragraphs; captions remain useful in text.
        if para.text.strip():
            clone_paragraph(doc, para)

    doc.add_page_break()
    doc.add_heading("TABLES", level=1)

    top_meta = meta.head(12).copy()
    p = doc.add_paragraph()
    p.add_run("Table 1. Top 12 genes ranked by meta-priority.").bold = True
    t1 = doc.add_table(rows=len(top_meta) + 1, cols=7)
    t1.style = "Table Grid"
    h1 = ["Meta rank", "Gene", "Evidence tier", "Support count", "Pooled effect", "95% CI", "I2"]
    for i, h in enumerate(h1):
        cell = t1.rows[0].cells[i]
        cell.text = h
        cell.paragraphs[0].runs[0].bold = True
        set_cell_shading(cell, "D9E2F3")
    for i, (_, row) in enumerate(top_meta.iterrows(), start=1):
        vals = [
            int(row["MetaRank"]), row["GeneSymbol"], row["EvidenceTier"], int(row["NominalSupportCount"]),
            f"{row['RandomEffect']:.2f}", f"{row['Lower95CI']:.2f} to {row['Upper95CI']:.2f}", f"{row['I2']:.1f}"
        ]
        for j, v in enumerate(vals):
            t1.rows[i].cells[j].text = str(v)

    p = doc.add_paragraph()
    p.add_run("Table 2. Translationally relevant targets with evidence tier and interpretation.").bold = True
    t2 = doc.add_table(rows=len(transl) + 1, cols=7)
    t2.style = "Table Grid"
    h2 = ["Meta rank", "Gene", "Pathway", "Evidence tier", "Support", "Pooled effect", "Interpretation"]
    interp = {
        "single-cohort": "Transcriptomic signal present but not recurrent",
        "mechanistic-only": "Mechanistically plausible but weak transcriptomic support",
        "cross-cohort": "Recurrent transcriptomic support",
    }
    for i, h in enumerate(h2):
        cell = t2.rows[0].cells[i]
        cell.text = h
        cell.paragraphs[0].runs[0].bold = True
        set_cell_shading(cell, "D9E2F3")
    for i, (_, row) in enumerate(transl.iterrows(), start=1):
        vals = [
            int(row["MetaRank"]), row["GeneSymbol"], row["Pathway"], row["EvidenceTier"],
            int(row["NominalSupportCount"]), f"{row['RandomEffect']:.2f}", interp[row["EvidenceTier"]]
        ]
        for j, v in enumerate(vals):
            t2.rows[i].cells[j].text = str(v)

    shortlist = revision_drugs[revision_drugs["GeneSymbol"].isin(["IL1B", "IL6", "TNF", "ANGPT2", "SERPINE1", "HMOX1", "F3", "VWF"])].copy()
    p = doc.add_paragraph()
    p.add_run("Table 3. Hypothesis-generating intervention shortlist under the final evidence framework.").bold = True
    t3 = doc.add_table(rows=len(shortlist) + 1, cols=5)
    t3.style = "Table Grid"
    h3 = ["Candidate", "Target", "Pathway", "Evidence tier", "Use in manuscript"]
    for i, h in enumerate(h3):
        cell = t3.rows[0].cells[i]
        cell.text = h
        cell.paragraphs[0].runs[0].bold = True
        set_cell_shading(cell, "D9E2F3")
    for i, (_, row) in enumerate(shortlist.iterrows(), start=1):
        vals = [row["Candidate"], row["GeneSymbol"], row["Pathway"], tier_lookup.get(row["GeneSymbol"], "n/a"), "Hypothesis-generating only"]
        for j, v in enumerate(vals):
            t3.rows[i].cells[j].text = str(v)

    out = MANUSCRIPT_DIR / "Manuscript_KFD_MJDYPV_Final_Blinded_TablesAfterRefs.docx"
    doc.save(out)
    return out


def build_figures_docx() -> Path:
    doc = Document()
    set_margins(doc)
    set_style(doc)
    h = doc.add_heading("", level=0)
    r = h.add_run("Figures")
    r.bold = True
    r.font.size = Pt(14)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("A Cross-Flaviviral Transcriptomic Evidence and Mechanistic Prioritization Framework for Host-Directed Therapy in Kyasanur Forest Disease")

    figures = [
        (FINAL_FIGS / "figure1_submission.png", "Figure 1. Discovery cohorts used in the analysis."),
        (FINAL_FIGS / "figure2_submission.png", "Figure 2. Meta-priority ranking after adding pooled effects and evidence tiers."),
        (FINAL_FIGS / "figure3_submission.png", "Figure 3. Pathway effect size versus heterogeneity."),
        (FINAL_FIGS / "figure4_submission.png", "Figure 4. Original composite ranking retained for comparison with the meta-analytic ranking."),
    ]
    for path, caption in figures:
        p = doc.add_paragraph()
        p.add_run(caption).bold = True
        doc.add_picture(str(path), width=Inches(6.0))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph()

    out = MANUSCRIPT_DIR / "FIGURES_KFD_MJDYPV_Final.docx"
    doc.save(out)
    return out


def main() -> None:
    tables_doc = build_tables_after_references_variant()
    figures_doc = build_figures_docx()
    print("Generated layout variants:")
    print(f" - {tables_doc.name}")
    print(f" - {figures_doc.name}")


if __name__ == "__main__":
    main()
