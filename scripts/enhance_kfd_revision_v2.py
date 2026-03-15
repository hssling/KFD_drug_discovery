"""Additive scientific-enhancement analysis for the KFD revision package.

This script does not modify existing revision assets. It creates a v2 layer
with target-level meta-analysis, uncertainty summaries, and companion figures.
"""

from __future__ import annotations

import math
from pathlib import Path

import matplotlib.pyplot as plt
import numpy as np
import pandas as pd
import seaborn as sns
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt
from scipy import stats


BASE_DIR = Path(__file__).resolve().parent.parent
REV_TABLES = BASE_DIR / "outputs" / "revision_tables"
V2_TABLES = BASE_DIR / "outputs" / "enhanced_v2_tables"
V2_FIGS = BASE_DIR / "outputs" / "enhanced_v2_figures"
MANUSCRIPTS = BASE_DIR / "manuscripts"

for directory in (V2_TABLES, V2_FIGS):
    directory.mkdir(parents=True, exist_ok=True)


def set_margins(doc: Document) -> None:
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)


def set_base_style(doc: Document, size: int = 11) -> None:
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(size)
    style.paragraph_format.line_spacing = 1.5


def set_cell_shading(cell, color: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), color)
    tc_pr.append(shd)


def z_from_pvalue(pvalue: float) -> float:
    if pd.isna(pvalue) or pvalue <= 0:
        return np.inf
    if pvalue >= 1:
        return 0.0
    return stats.norm.isf(pvalue / 2.0)


def approximate_se(log2fc: float, pvalue: float) -> float:
    z = z_from_pvalue(pvalue)
    if not np.isfinite(z) or z == 0 or pd.isna(log2fc):
        return np.nan
    return abs(log2fc) / z


def dersimonian_laird_meta(effects: np.ndarray, ses: np.ndarray) -> dict[str, float]:
    mask = np.isfinite(effects) & np.isfinite(ses) & (ses > 0)
    effects = effects[mask]
    ses = ses[mask]
    k = len(effects)
    if k == 0:
        return {
            "Studies": 0,
            "FixedEffect": np.nan,
            "FixedSE": np.nan,
            "RandomEffect": np.nan,
            "RandomSE": np.nan,
            "Tau2": np.nan,
            "Q": np.nan,
            "I2": np.nan,
            "PooledPValue": np.nan,
        }
    w = 1 / (ses ** 2)
    fixed_effect = float(np.sum(w * effects) / np.sum(w))
    fixed_se = float(np.sqrt(1 / np.sum(w)))
    if k == 1:
        pooled_p = 2 * stats.norm.sf(abs(fixed_effect / fixed_se)) if fixed_se > 0 else np.nan
        return {
            "Studies": 1,
            "FixedEffect": fixed_effect,
            "FixedSE": fixed_se,
            "RandomEffect": fixed_effect,
            "RandomSE": fixed_se,
            "Tau2": 0.0,
            "Q": 0.0,
            "I2": 0.0,
            "PooledPValue": pooled_p,
        }
    q = float(np.sum(w * (effects - fixed_effect) ** 2))
    c = float(np.sum(w) - (np.sum(w ** 2) / np.sum(w)))
    tau2 = max((q - (k - 1)) / c, 0.0) if c > 0 else 0.0
    w_re = 1 / (ses ** 2 + tau2)
    random_effect = float(np.sum(w_re * effects) / np.sum(w_re))
    random_se = float(np.sqrt(1 / np.sum(w_re)))
    pooled_p = 2 * stats.norm.sf(abs(random_effect / random_se)) if random_se > 0 else np.nan
    i2 = max((q - (k - 1)) / q, 0.0) * 100 if q > 0 else 0.0
    return {
        "Studies": k,
        "FixedEffect": fixed_effect,
        "FixedSE": fixed_se,
        "RandomEffect": random_effect,
        "RandomSE": random_se,
        "Tau2": tau2,
        "Q": q,
        "I2": i2,
        "PooledPValue": pooled_p,
    }


def build_meta_table() -> pd.DataFrame:
    panel = pd.read_csv(BASE_DIR / "data" / "gene_signature.csv").rename(columns={"Symbol": "GeneSymbol"})
    deg_tables = {
        "GSE18090": pd.read_csv(REV_TABLES / "GSE18090_deg_results.csv"),
        "GSE51808": pd.read_csv(REV_TABLES / "GSE51808_deg_results.csv"),
        "GSE43777": pd.read_csv(REV_TABLES / "GSE43777_deg_results.csv"),
    }

    records = []
    for _, panel_row in panel.iterrows():
        per_study = []
        directions = []
        nominal_support = 0
        for study_name, df in deg_tables.items():
            hit = df[df["GeneSymbol"] == panel_row["GeneSymbol"]]
            if hit.empty:
                continue
            row = hit.iloc[0]
            effect = float(row["log2FC"])
            pvalue = float(row["pvalue"])
            se = approximate_se(effect, pvalue)
            per_study.append((study_name, effect, se, pvalue))
            directions.append(np.sign(effect))
            if pvalue <= 0.05 and abs(effect) >= 0.30:
                nominal_support += 1

        meta = dersimonian_laird_meta(
            np.array([value[1] for value in per_study], dtype=float),
            np.array([value[2] for value in per_study], dtype=float),
        )
        direction_labels = ["up" if d > 0 else "down" if d < 0 else "flat" for d in directions]
        direction_concordance = len(set(direction_labels)) == 1 if direction_labels else False
        pooled_direction = (
            "up" if meta["RandomEffect"] > 0 else "down" if meta["RandomEffect"] < 0 else "flat"
        )
        evidence_tier = (
            "cross-cohort"
            if nominal_support >= 2 and meta["PooledPValue"] <= 0.05
            else "single-cohort"
            if nominal_support == 1
            else "mechanistic-only"
        )

        records.append(
            {
                "GeneSymbol": panel_row["GeneSymbol"],
                "GeneName": panel_row["Gene"],
                "Pathway": panel_row["Pathway"],
                "PhaseRelevance": panel_row["Phase_Relevance"],
                "Druggability": panel_row["Druggability"],
                "Studies": meta["Studies"],
                "NominalSupportCount": nominal_support,
                "DirectionConcordant": direction_concordance,
                "PooledDirection": pooled_direction,
                "RandomEffect": meta["RandomEffect"],
                "RandomSE": meta["RandomSE"],
                "Lower95CI": meta["RandomEffect"] - 1.96 * meta["RandomSE"] if pd.notna(meta["RandomSE"]) else np.nan,
                "Upper95CI": meta["RandomEffect"] + 1.96 * meta["RandomSE"] if pd.notna(meta["RandomSE"]) else np.nan,
                "PooledPValue": meta["PooledPValue"],
                "I2": meta["I2"],
                "Tau2": meta["Tau2"],
                "EvidenceTier": evidence_tier,
                "PerStudyEffects": "; ".join(
                    f"{study}:{effect:.2f}, p={pvalue:.3g}" for study, effect, _, pvalue in per_study
                ),
            }
        )

    meta_df = pd.DataFrame(records)
    meta_df["AbsRandomEffect"] = meta_df["RandomEffect"].abs()
    meta_df["MetaPriority"] = (
        0.45 * np.clip(meta_df["NominalSupportCount"] / 3.0, 0, 1)
        + 0.35 * np.clip(meta_df["AbsRandomEffect"] / 1.0, 0, 1)
        + 0.20 * (1 - np.clip(meta_df["I2"].fillna(100) / 100.0, 0, 1))
    )
    meta_df = meta_df.sort_values(
        ["EvidenceTier", "MetaPriority", "PooledPValue"],
        ascending=[True, False, True],
        key=lambda s: s.map({"cross-cohort": 0, "single-cohort": 1, "mechanistic-only": 2}) if s.name == "EvidenceTier" else s,
    ).reset_index(drop=True)
    meta_df["MetaRank"] = np.arange(1, len(meta_df) + 1)
    return meta_df


def make_figures(meta_df: pd.DataFrame) -> None:
    sns.set_theme(style="whitegrid")

    top = meta_df.head(20).copy()
    fig, ax = plt.subplots(figsize=(9, 7))
    sns.barplot(data=top, y="GeneSymbol", x="MetaPriority", hue="EvidenceTier", dodge=False, ax=ax)
    ax.set_title("Enhanced v2 Figure 1. Meta-analytic target prioritization")
    ax.set_xlabel("Meta-priority score")
    ax.set_ylabel("Gene")
    fig.tight_layout()
    fig.savefig(V2_FIGS / "figure_v2_meta_priority.png", dpi=300)
    plt.close(fig)

    pathway = (
        meta_df.groupby("Pathway")
        .agg(
            MeanAbsEffect=("AbsRandomEffect", "mean"),
            MeanI2=("I2", "mean"),
            CrossCohort=("EvidenceTier", lambda x: int((x == "cross-cohort").sum())),
        )
        .reset_index()
    )
    fig, ax = plt.subplots(figsize=(8, 5))
    sns.scatterplot(data=pathway, x="MeanAbsEffect", y="MeanI2", size="CrossCohort", hue="Pathway", ax=ax, sizes=(50, 300))
    ax.set_title("Enhanced v2 Figure 2. Pathway effect size versus heterogeneity")
    ax.set_xlabel("Mean absolute pooled effect")
    ax.set_ylabel("Mean I-squared")
    fig.tight_layout()
    fig.savefig(V2_FIGS / "figure_v2_pathway_heterogeneity.png", dpi=300)
    plt.close(fig)


def write_tables(meta_df: pd.DataFrame) -> None:
    meta_df.to_csv(V2_TABLES / "kfd_enhanced_v2_meta_targets.csv", index=False)
    summary = (
        meta_df.groupby(["Pathway", "EvidenceTier"])
        .size()
        .reset_index(name="Targets")
        .sort_values(["Pathway", "EvidenceTier"])
    )
    summary.to_csv(V2_TABLES / "kfd_enhanced_v2_evidence_summary.csv", index=False)

    translational = meta_df.loc[
        meta_df["GeneSymbol"].isin(["IL1B", "IL6", "TNF", "ANGPT2", "F3", "VWF", "SERPINE1", "HMOX1", "BDNF"]),
        [
            "MetaRank", "GeneSymbol", "Pathway", "NominalSupportCount", "RandomEffect",
            "Lower95CI", "Upper95CI", "I2", "EvidenceTier", "PerStudyEffects"
        ],
    ].sort_values("MetaRank")
    translational.to_csv(V2_TABLES / "kfd_enhanced_v2_translational_targets.csv", index=False)


def build_memo(meta_df: pd.DataFrame) -> Path:
    summary = pd.read_csv(V2_TABLES / "kfd_enhanced_v2_evidence_summary.csv")
    translational = pd.read_csv(V2_TABLES / "kfd_enhanced_v2_translational_targets.csv")
    cross = meta_df[meta_df["EvidenceTier"] == "cross-cohort"]
    single = meta_df[meta_df["EvidenceTier"] == "single-cohort"]

    doc = Document()
    set_margins(doc)
    set_base_style(doc)

    title = doc.add_heading("", level=0)
    run = title.add_run("KFD Scientific Enhancement Memo v2")
    run.bold = True
    run.font.size = Pt(14)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph(
        "Purpose: add a stronger statistical layer to the current revision without changing existing submission assets."
    )

    for paragraph in [
        f"Meta-analysis was added for the prespecified 50-gene panel using the existing cohort-level severe-versus-non-severe effect estimates from GSE18090, GSE51808, and GSE43777. Standard errors were approximated from two-sided P values and log2 fold-changes, then pooled with a DerSimonian-Laird random-effects model.",
        f"Only {len(cross)} genes met a cross-cohort evidence tier, whereas {len(single)} genes showed single-cohort nominal support. This confirms that the strongest evidence in the current dataset base is concentrated in a limited subset of targets, while many endothelial/coagulation genes remain mechanistic-priority hypotheses rather than recurrent transcriptomic findings.",
        "This v2 layer strengthens rigor in three ways: it provides pooled effects with confidence intervals, quantifies heterogeneity, and separates cross-cohort versus single-cohort evidence. It therefore supports more precise wording around which conclusions are well supported and which remain exploratory.",
    ]:
        doc.add_paragraph(paragraph)

    doc.add_heading("Table 1. Evidence-tier summary", level=1)
    table1 = doc.add_table(rows=len(summary) + 1, cols=len(summary.columns))
    table1.style = "Table Grid"
    for idx, col in enumerate(summary.columns):
        cell = table1.rows[0].cells[idx]
        cell.text = str(col)
        cell.paragraphs[0].runs[0].bold = True
        set_cell_shading(cell, "D9E2F3")
    for r_idx, (_, row) in enumerate(summary.iterrows(), start=1):
        for c_idx, col in enumerate(summary.columns):
            table1.rows[r_idx].cells[c_idx].text = str(row[col])

    doc.add_heading("Table 2. Translationally important targets with uncertainty estimates", level=1)
    table2 = doc.add_table(rows=len(translational) + 1, cols=7)
    table2.style = "Table Grid"
    headers = ["Meta rank", "Gene", "Pathway", "Support", "Pooled effect (95% CI)", "I2", "Evidence tier"]
    for idx, header in enumerate(headers):
        cell = table2.rows[0].cells[idx]
        cell.text = header
        cell.paragraphs[0].runs[0].bold = True
        set_cell_shading(cell, "D9E2F3")
    for r_idx, (_, row) in enumerate(translational.iterrows(), start=1):
        values = [
            int(row["MetaRank"]),
            row["GeneSymbol"],
            row["Pathway"],
            int(row["NominalSupportCount"]),
            f"{row['RandomEffect']:.2f} ({row['Lower95CI']:.2f} to {row['Upper95CI']:.2f})",
            f"{row['I2']:.1f}",
            row["EvidenceTier"],
        ]
        for c_idx, value in enumerate(values):
            table2.rows[r_idx].cells[c_idx].text = str(value)

    doc.add_heading("Recommended wording upgrade", level=1)
    for bullet in [
        "Inflammatory genes have the strongest transcriptomic support across cohorts.",
        "Endothelial and coagulation genes should be described as clinically motivated mechanistic priorities with limited cross-cohort recurrence in current public data.",
        "The drug shortlist should remain explicitly hypothesis-generating and staged for preclinical validation.",
    ]:
        doc.add_paragraph(bullet, style="List Bullet")

    out_path = MANUSCRIPTS / "KFD_Scientific_Enhancement_Memo_v2.docx"
    doc.save(out_path)
    return out_path


def main() -> None:
    meta_df = build_meta_table()
    write_tables(meta_df)
    make_figures(meta_df)
    memo_path = build_memo(meta_df)
    print("Generated additive v2 enhancement package:")
    print(f" - {memo_path.name}")
    print(f" - {V2_TABLES / 'kfd_enhanced_v2_meta_targets.csv'}")
    print(f" - {V2_FIGS / 'figure_v2_meta_priority.png'}")


if __name__ == "__main__":
    main()
