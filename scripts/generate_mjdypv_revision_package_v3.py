"""Generate an additive v3 manuscript package incorporating v2 meta-analysis."""

from __future__ import annotations

import re
from datetime import datetime
from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt


BASE_DIR = Path(__file__).resolve().parent.parent
MANUSCRIPT_DIR = BASE_DIR / "manuscripts"
REV_TABLES = BASE_DIR / "outputs" / "revision_tables"
V2_TABLES = BASE_DIR / "outputs" / "enhanced_v2_tables"
REV_FIGS = BASE_DIR / "outputs" / "revision_figures"
V2_FIGS = BASE_DIR / "outputs" / "enhanced_v2_figures"

TITLE = "A Cross-Flaviviral Transcriptomic Evidence and Mechanistic Prioritization Framework for Host-Directed Therapy in Kyasanur Forest Disease"
RUNNING_TITLE = "Evidence-Based HDT Framework for KFD"


def set_margins(doc: Document) -> None:
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)


def set_base_style(doc: Document, size: int = 12) -> None:
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(size)
    style.paragraph_format.line_spacing = 2.0


def set_cell_shading(cell, color: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), color)
    tc_pr.append(shd)


def add_cited_paragraph(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph()
    parts = re.split(r"(\[\d+(?:[-,]\d+)*\])", text)
    for part in parts:
        if part.startswith("[") and part.endswith("]"):
            run = paragraph.add_run(part)
            run.font.superscript = True
        else:
            paragraph.add_run(part)


def references() -> list[str]:
    return [
        "Work TH, Trapido H, Murthy DP, Rao RL, Bhatt PN, Kulkarni KG. Kyasanur forest disease. III. A preliminary report on the nature of the infection and clinical manifestations in man. Indian J Med Sci 1957;11:619-45.",
        "Pattnaik P. Kyasanur forest disease: an epidemiological view in India. Rev Med Virol 2006;16:151-65.",
        "Murhekar MV, Kasabi GS, Mehendale SM, Mourya DT, Yadav PD, Tandale BV, et al. On the transmission pattern of Kyasanur Forest Disease (KFD) in India. Infect Dis Poverty 2015;4:37.",
        "Holbrook MR. Kyasanur forest disease. Antiviral Res 2012;96:353-62.",
        "Kasabi GS, Murhekar MV, Sandhya VK, Raghunandan R, Kiran SK, Channabasappa GH, et al. Coverage and effectiveness of Kyasanur forest disease vaccine in Karnataka, South India, 2005-10. PLoS Negl Trop Dis 2013;7:e2025.",
        "Kaufmann SHE, Dorhoi A, Hotchkiss RS, Bartenschlager R. Host-directed therapies for bacterial and viral infections. Nat Rev Drug Discov 2018;17:35-56.",
        "Zumla A, Rao M, Wallis RS, Kaufmann SH, Rustomjee R, Mwaba P, et al. Host-directed therapies for infectious diseases: current status, recent progress, and future prospects. Lancet Infect Dis 2016;16:e47-63.",
        "Barrett T, Wilhite SE, Ledoux P, Evangelista C, Kim IF, Tomashevsky M, et al. NCBI GEO: archive for functional genomics data sets-update. Nucleic Acids Res 2013;41:D991-5.",
        "Nascimento EJ, Braga-Neto U, Calzavara-Silva CE, Gomes AL, Abath FG, Brito CA, et al. Gene expression profiling during early acute febrile stage of dengue infection can predict the disease outcome. PLoS One 2009;4:e7892.",
        "Sun P, Garcia J, Comach G, Vahey MT, Wang Z, Forshey BM, et al. Sequential waves of gene expression in patients with clinically defined dengue illnesses reveal subtle disease phases and predict disease severity. PLoS Negl Trop Dis 2013;7:e2298.",
        "Kwissa M, Nakaya HI, Onlamoon N, Wrammert J, Villinger F, Perng GC, et al. Dengue virus infection induces expansion of a CD14+CD16+ monocyte population that stimulates plasmablast differentiation. Cell Host Microbe 2014;16:115-27.",
        "Gillespie M, Jassal B, Stephan R, Milacic M, Rothfels K, Senff-Ribeiro A, et al. The Reactome pathway knowledgebase 2022. Nucleic Acids Res 2022;50:D687-92.",
        "Zdrazil B, Felix E, Hunter F, Manners EJ, Blackshaw J, Corbett S, et al. The ChEMBL Database in 2023: a drug discovery platform spanning multiple bioactivity data types and time periods. Nucleic Acids Res 2024;52:D1180-92.",
        "Simmons CP, Farrar JJ, Nguyen vV, Wills B. Dengue. N Engl J Med 2012;366:1423-32.",
        "Modhiran N, Watterson D, Muller DA, Panetta AK, Sester DP, Liu L, et al. Dengue virus NS1 protein activates cells via Toll-like receptor 4 and disrupts endothelial cell monolayer integrity. Sci Transl Med 2015;7:304ra142.",
        "Bray M. Pathogenesis of viral hemorrhagic fever. Curr Opin Immunol 2005;17:399-403.",
        "Yacoub S, Wills B. Predicting outcome from dengue. BMC Med 2014;12:147.",
        "Roberts I, Shakur H, Coats T, Hunt B, Balogun E, Barnetson L, et al. The CRASH-2 trial: a randomised controlled trial and economic evaluation of the effects of tranexamic acid on death, vascular occlusive events and transfusion requirement in bleeding trauma patients. Health Technol Assess 2013;17:1-79.",
        "Kiran SK, Padamashree S, Jayashree K, Srinath S. Health-care-seeking behaviour among Kyasanur forest disease patients in Shivamogga district, Karnataka: a cross-sectional study. Indian J Community Med 2021;46:486-9.",
    ]


def build_package() -> list[Path]:
    meta = pd.read_csv(V2_TABLES / "kfd_enhanced_v2_meta_targets.csv")
    transl = pd.read_csv(V2_TABLES / "kfd_enhanced_v2_translational_targets.csv")
    revision_targets = pd.read_csv(REV_TABLES / "kfd_revision_targets.csv")
    cohorts = pd.read_csv(REV_TABLES / "cohort_summary.csv")

    severe = int(cohorts["SevereSamples"].sum())
    non_severe = int(cohorts["NonSevereSamples"].sum())
    total = severe + non_severe
    single = int((meta["EvidenceTier"] == "single-cohort").sum())
    mech = int((meta["EvidenceTier"] == "mechanistic-only").sum())

    manuscript = Document()
    set_margins(manuscript)
    set_base_style(manuscript)

    title = manuscript.add_heading("", level=0)
    run = title.add_run(TITLE)
    run.bold = True
    run.font.size = Pt(14)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = manuscript.add_paragraph()
    p.add_run("Running Title: ").bold = True
    p.add_run(RUNNING_TITLE)
    manuscript.add_page_break()

    manuscript.add_heading("ABSTRACT", level=1)
    abstract = [
        ("Background:", "Kyasanur Forest Disease lacks disease-specific transcriptomic datasets and specific therapy, forcing host-directed therapeutic work to rely on proxy flaviviral data.[1-5]"),
        ("Objectives:", "To strengthen a transcriptomic prioritization framework for KFD by adding pooled effect estimates, confidence intervals, and heterogeneity metrics to the existing mechanistic ranking."),
        ("Materials and Methods:", f"We reanalyzed three public dengue severity cohorts from GEO (GSE18090, GSE43777, GSE51808; {total} acute samples, {severe} severe and {non_severe} non-severe). A prespecified 50-gene host-response panel was scored as in the revised submission, then supplemented with random-effects meta-analysis, 95% confidence intervals, and evidence-tier classification."),
        ("Results:", f"The strongest transcriptomic evidence remained inflammatory. No gene reached a strict cross-cohort evidence tier; {single} genes had single-cohort nominal support and {mech} remained mechanistic-only. The best-supported pooled signals included CXCL10, IL6, STAT1, and IFNB1, whereas endothelial/coagulation genes such as ANGPT2, F3, and VWF remained biologically plausible but weakly supported transcriptomically."),
        ("Conclusions:", "The v3 framing improves validity by separating transcriptomic evidence from mechanistic plausibility. Inflammatory targets are the best-supported findings in current public data, while endothelial and coagulation pathways remain clinically motivated hypotheses for KFD that require KFD-specific validation before therapeutic inference.[6-18]"),
    ]
    for label, text in abstract:
        p = manuscript.add_paragraph()
        p.add_run(label).bold = True
        p.add_run(" " + text)
    p = manuscript.add_paragraph()
    p.add_run("Keywords: ").bold = True
    p.add_run("Kyasanur Forest Disease; transcriptomics; meta-analysis; host-directed therapy; uncertainty; flavivirus")
    manuscript.add_page_break()

    manuscript.add_heading("INTRODUCTION", level=1)
    for text in [
        "Kyasanur Forest Disease (KFD) remains a clinically important tick-borne flaviviral hemorrhagic fever in southern India, with no specific antiviral therapy and incomplete vaccine protection.[1-5]",
        "The revised manuscript established a transparent transcriptomic prioritization framework, but the public data base remained indirect because KFD-specific blood transcriptomes were unavailable. The current v3 version adds an explicit uncertainty layer so that transcriptomic support and mechanistic prioritization are no longer conflated.",
        "This distinction matters because host-directed therapy papers can become clinically misleading when plausible vascular or coagulation mechanisms are presented as if they were strongly supported by recurrent transcriptomic evidence. The purpose of v3 is to make that evidentiary boundary explicit while preserving the clinically relevant mechanistic hypotheses.[6,7,14-18]",
    ]:
        add_cited_paragraph(manuscript, text)

    manuscript.add_heading("MATERIALS AND METHODS", level=1)
    for text in [
        f"We retained the same three discovery cohorts used in the revised submission: GSE18090, GSE43777, and GSE51808, totaling {total} acute samples ({severe} severe and {non_severe} non-severe). Dataset preprocessing, probe collapsing, and within-cohort severe-versus-non-severe contrasts were unchanged.",
        "The v3 enhancement added a random-effects meta-analysis for each of the 50 prespecified panel genes. Standard errors were approximated from cohort-level log2 fold-change and two-sided P values, after which pooled random-effects estimates, 95% confidence intervals, and heterogeneity statistics (I-squared and tau-squared) were calculated.",
        "Each gene was assigned an evidence tier: cross-cohort if at least two cohorts showed same-direction nominal support and the pooled evidence remained statistically convincing; single-cohort if one cohort showed nominal support; and mechanistic-only if prioritization depended primarily on pathway or tractability considerations rather than transcriptomic recurrence.",
    ]:
        manuscript.add_paragraph(text)

    manuscript.add_heading("RESULTS", level=1)
    for text in [
        f"No gene in the 50-gene panel reached the strict cross-cohort evidence tier. {single} genes had single-cohort nominal support and {mech} were classified as mechanistic-only. This result shows that the current public data base is better suited to ranking plausible host-response hypotheses than to establishing robust cross-flaviviral consensus biomarkers.",
        "The highest meta-priority genes were CXCL10, IL6, STAT1, IFNB1, and TEK. Among these, the most convincing pooled inflammatory signal was IL6, although even that estimate retained moderate heterogeneity and a confidence interval that crossed the null. CXCL10 ranked first by meta-priority because of effect size, but it showed substantial heterogeneity and direction inconsistency across cohorts.",
        "Translationally important endothelial and coagulation genes remained weaker on the transcriptomic evidence axis. ANGPT2, F3, and VWF stayed biologically relevant for KFD pathophysiology, but in v3 they are explicitly presented as mechanistic-priority targets with limited direct transcriptomic support in the public datasets.",
    ]:
        manuscript.add_paragraph(text)

    p = manuscript.add_paragraph()
    p.add_run("Table 1. Top 12 genes ranked by v2 meta-priority.").bold = True
    table1 = manuscript.add_table(rows=13, cols=7)
    table1.style = "Table Grid"
    headers = ["Meta rank", "Gene", "Evidence tier", "Support count", "Pooled effect", "95% CI", "I2"]
    for i, h in enumerate(headers):
        cell = table1.rows[0].cells[i]
        cell.text = h
        cell.paragraphs[0].runs[0].bold = True
        set_cell_shading(cell, "D9E2F3")
    for i, (_, row) in enumerate(meta.head(12).iterrows(), start=1):
        vals = [
            int(row["MetaRank"]),
            row["GeneSymbol"],
            row["EvidenceTier"],
            int(row["NominalSupportCount"]),
            f"{row['RandomEffect']:.2f}",
            f"{row['Lower95CI']:.2f} to {row['Upper95CI']:.2f}",
            f"{row['I2']:.1f}",
        ]
        for j, v in enumerate(vals):
            table1.rows[i].cells[j].text = str(v)

    p = manuscript.add_paragraph()
    p.add_run("Table 2. Translationally relevant inflammatory, endothelial, coagulation, oxidative, and neurological targets.").bold = True
    table2 = manuscript.add_table(rows=len(transl) + 1, cols=7)
    table2.style = "Table Grid"
    headers2 = ["Meta rank", "Gene", "Pathway", "Evidence tier", "Support", "Pooled effect", "Interpretation"]
    for i, h in enumerate(headers2):
        cell = table2.rows[0].cells[i]
        cell.text = h
        cell.paragraphs[0].runs[0].bold = True
        set_cell_shading(cell, "D9E2F3")
    interp_map = {
        "single-cohort": "Transcriptomic signal present but not recurrent",
        "mechanistic-only": "Mechanistically plausible but weak transcriptomic support",
        "cross-cohort": "Recurrent transcriptomic support",
    }
    for i, (_, row) in enumerate(transl.iterrows(), start=1):
        vals = [
            int(row["MetaRank"]),
            row["GeneSymbol"],
            row["Pathway"],
            row["EvidenceTier"],
            int(row["NominalSupportCount"]),
            f"{row['RandomEffect']:.2f}",
            interp_map[row["EvidenceTier"]],
        ]
        for j, v in enumerate(vals):
            table2.rows[i].cells[j].text = str(v)

    for filename, caption in [
        ("figure_v2_meta_priority.png", "Figure 1. Meta-priority ranking after adding pooled effects and evidence tiers."),
        ("figure_v2_pathway_heterogeneity.png", "Figure 2. Pathway effect size versus heterogeneity, showing how uncertainty varies across modules."),
        ("figure2_target_ranking.png", "Figure 3. Original composite-ranking view retained for comparison with the meta-analytic ranking."),
    ]:
        p = manuscript.add_paragraph()
        p.add_run(caption).bold = True
        source = V2_FIGS / filename if filename.startswith("figure_v2") else REV_FIGS / filename
        manuscript.add_picture(str(source), width=Inches(5.7))
        manuscript.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    manuscript.add_heading("DISCUSSION", level=1)
    for text in [
        "The main scientific value of the v3 enhancement is not that it dramatically improves the therapeutic shortlist, but that it improves interpretability. The added meta-analysis shows that inflammatory targets remain the clearest transcriptomic findings, whereas many vascular/coagulation targets derive their priority mainly from KFD pathophysiology and tractability, not recurrent transcriptomic evidence.",
        "This makes the manuscript more robust because it aligns conclusions with data strength. A reader can now distinguish between evidence-based inflammatory signals and mechanistically prioritized endothelial/coagulation hypotheses. That is a more honest and scientifically useful basis for future translational work.",
        "The price of this stronger framing is that the translational claims become narrower. The current public proxy datasets are insufficient to validate a KFD-specific host-response signature, and they do not justify therapeutic recommendation. They do, however, justify a focused validation agenda centered on inflammatory markers plus selected vascular/coagulation candidates.",
    ]:
        manuscript.add_paragraph(text)

    manuscript.add_heading("LIMITATIONS", level=2)
    manuscript.add_paragraph(
        "The v3 version still depends entirely on non-KFD public blood transcriptomes. The meta-analysis uses approximate standard errors reconstructed from published effect sizes and P values rather than raw per-sample variance models. Evidence tiers are therefore useful for ranking uncertainty, but they do not replace true KFD-specific validation."
    )

    manuscript.add_heading("CONCLUSIONS", level=1)
    manuscript.add_paragraph(
        "The v3 manuscript is the most scientifically conservative and statistically explicit version of the current project. It supports inflammatory targets as the best transcriptomic findings in the available public data and retains endothelial/coagulation pathways as mechanistic KFD hypotheses that require disease-specific validation."
    )

    manuscript.add_page_break()
    manuscript.add_heading("REFERENCES", level=1)
    for i, ref in enumerate(references(), start=1):
        p = manuscript.add_paragraph()
        p.add_run(f"{i}. ").bold = True
        p.add_run(ref)
        p.paragraph_format.first_line_indent = Inches(-0.25)
        p.paragraph_format.left_indent = Inches(0.25)

    blinded = MANUSCRIPT_DIR / "Manuscript_KFD_MJDYPV_v3_Blinded.docx"
    manuscript.save(blinded)

    titlepage = Document()
    set_margins(titlepage)
    set_base_style(titlepage)
    for label, value in [
        ("Article Type", "Original Article"),
        ("Title", TITLE),
        ("Running Title", RUNNING_TITLE),
        ("Author", "Siddalingaiah H S"),
        ("Affiliation", "Department of Community Medicine, Shridevi Institute of Medical Sciences and Research Hospital, Tumkur, Karnataka, India"),
        ("Corresponding Author", "Dr. Siddalingaiah H S, MD; hssling@yahoo.com; +91-8941087719; ORCID 0000-0002-4771-8285"),
        ("Version Note", "v3 package integrates additive meta-analysis/uncertainty layer without altering the original revised submission"),
    ]:
        p = titlepage.add_paragraph()
        p.add_run(f"{label}: ").bold = True
        p.add_run(value)
    title_path = MANUSCRIPT_DIR / "TitlePage_KFD_MJDYPV_v3.docx"
    titlepage.save(title_path)

    note = Document()
    set_margins(note)
    set_base_style(note, 11)
    h = note.add_heading("", level=0)
    r = h.add_run("v3 Statistical Framing Note")
    r.bold = True
    r.font.size = Pt(14)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for text in [
        "This v3 package is an additive manuscript variant that incorporates the v2 meta-analysis outputs.",
        "It is intended for a stronger future submission or major revision, not as a replacement for the currently prepared revised submission package.",
        f"Key shift: 0 genes met the strict cross-cohort tier; {single} genes were single-cohort; {mech} were mechanistic-only. Accordingly, v3 narrows the strongest claims to inflammatory signaling and reframes vascular/coagulation findings as mechanistic hypotheses.",
    ]:
        note.add_paragraph(text)
    note_path = MANUSCRIPT_DIR / "KFD_v3_Statistical_Framing_Note.docx"
    note.save(note_path)

    return [blinded, title_path, note_path]


def main() -> None:
    files = build_package()
    print("Generated v3 manuscript package:")
    for file in files:
        print(f" - {file.name}")


if __name__ == "__main__":
    main()
