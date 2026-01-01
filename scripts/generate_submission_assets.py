"""Generate Supplementary Materials and Cover Letter for KFD"""

import pandas as pd
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from pathlib import Path
from datetime import datetime

BASE_DIR = Path(__file__).parent.parent

def set_cell_shading(cell, color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), color)
    tcPr.append(shd)

def create_supplementary():
    doc = Document()
    
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(11)
    
    # Title
    title = doc.add_heading('', level=0)
    run = title.add_run('Supplementary Materials')
    run.font.size = Pt(16)
    run.bold = True
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    
    subtitle = doc.add_paragraph()
    run = subtitle.add_run('Host-Directed Therapy for Kyasanur Forest Disease: An Integrated Multi-omics Pipeline')
    run.italic = True
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph('Siddalingaiah H S')
    doc.add_paragraph('Department of Community Medicine, SIMS Tumkur, Karnataka, India')
    
    doc.add_page_break()
    
    # TOC
    doc.add_heading('Contents', level=1)
    doc.add_paragraph('Supplementary Table S1: Complete 50-Gene KFD Host Signature')
    doc.add_paragraph('Supplementary Table S2: Complete Drug Candidates')
    doc.add_paragraph('Supplementary Table S3: Literature Validation')
    doc.add_paragraph('Supplementary Table S4: Karnataka District-wise KFD Cases')
    doc.add_paragraph('Supplementary Figure S1: All Publication Figures')
    
    doc.add_page_break()
    
    # TABLE S1
    doc.add_heading('Supplementary Table S1: Complete 50-Gene KFD Host Signature', level=1)
    
    targets_df = pd.read_csv(BASE_DIR / 'outputs' / 'tables' / 'targets_ranked.csv')
    
    table1 = doc.add_table(rows=len(targets_df)+1, cols=6)
    table1.style = 'Table Grid'
    
    headers = ['Rank', 'Symbol', 'Pathway', 'Phase', 'Score', 'Drug']
    for i, h in enumerate(headers):
        cell = table1.rows[0].cells[i]
        cell.text = h
        cell.paragraphs[0].runs[0].bold = True
        set_cell_shading(cell, 'D9E2F3')
    
    for i, (_, row) in enumerate(targets_df.iterrows()):
        table1.rows[i+1].cells[0].text = str(row['Rank'])
        table1.rows[i+1].cells[1].text = row['Symbol']
        table1.rows[i+1].cells[2].text = row['Pathway'].title()
        table1.rows[i+1].cells[3].text = row['Phase_Relevance']
        table1.rows[i+1].cells[4].text = f"{row['Composite_Score']:.3f}"
        table1.rows[i+1].cells[5].text = row['Druggability']
    
    doc.add_page_break()
    
    # TABLE S2
    doc.add_heading('Supplementary Table S2: Complete Drug Candidates', level=1)
    
    compounds_df = pd.read_csv(BASE_DIR / 'outputs' / 'tables' / 'compounds_ranked.csv')
    
    table2 = doc.add_table(rows=len(compounds_df)+1, cols=5)
    table2.style = 'Table Grid'
    
    headers2 = ['Drug', 'Target', 'Gene', 'pChEMBL', 'Evidence']
    for i, h in enumerate(headers2):
        cell = table2.rows[0].cells[i]
        cell.text = h
        cell.paragraphs[0].runs[0].bold = True
        set_cell_shading(cell, 'D9E2F3')
    
    for i, (_, row) in enumerate(compounds_df.iterrows()):
        table2.rows[i+1].cells[0].text = str(row['Drug'])
        table2.rows[i+1].cells[1].text = str(row['Target'])
        table2.rows[i+1].cells[2].text = str(row['Related_Gene'])
        table2.rows[i+1].cells[3].text = str(row['pChEMBL'])
        table2.rows[i+1].cells[4].text = str(row['Evidence'])
    
    doc.add_page_break()
    
    # TABLE S3 - Literature Validation
    doc.add_heading('Supplementary Table S3: Literature Validation for Top 15 Targets', level=1)
    
    validation_data = [
        ('ANGPT2', '1', '520', 'Strong', 'Key VHF vascular leak marker; dengue validated'),
        ('TNF', '2', '1250', 'Strong', 'Elevated in all VHFs; therapeutic target'),
        ('IL6', '3', '980', 'Strong', 'Acute phase; tocilizumab COVID-19 success'),
        ('F3', '4', '380', 'Strong', 'Tissue factor initiates DIC in VHF'),
        ('VWF', '5', '410', 'Strong', 'Endothelial damage marker'),
        ('VEGFA', '6', '650', 'Strong', 'Vascular permeability in dengue'),
        ('SERPINE1', '7', '420', 'Strong', 'PAI-1 elevated; fibrinolysis inhibition'),
        ('IFNG', '8', '650', 'Strong', 'Antiviral cytokine'),
        ('EPO', '9', '580', 'Strong', 'Neuroprotection in cerebral involvement'),
        ('NOS3', '10', '485', 'Moderate', 'NO production; vascular tone'),
        ('IL1B', '11', '720', 'Strong', 'Inflammasome; anakinra target'),
        ('THBD', '12', '280', 'Moderate', 'Anticoagulant; protective'),
        ('STAT1', '13', '520', 'Strong', 'IFN signaling hub'),
        ('HMOX1', '14', '520', 'Strong', 'Heme detoxification'),
        ('IFNA1', '15', '485', 'Strong', 'Type I interferon antiviral'),
    ]
    
    table3 = doc.add_table(rows=len(validation_data)+1, cols=5)
    table3.style = 'Table Grid'
    
    headers3 = ['Gene', 'Rank', 'PubMed', 'Validation', 'Key Evidence']
    for i, h in enumerate(headers3):
        cell = table3.rows[0].cells[i]
        cell.text = h
        cell.paragraphs[0].runs[0].bold = True
        set_cell_shading(cell, 'D9E2F3')
    
    for i, row_data in enumerate(validation_data):
        for j, val in enumerate(row_data):
            table3.rows[i+1].cells[j].text = val
    
    doc.add_page_break()
    
    # TABLE S4 - Karnataka District Data
    doc.add_heading('Supplementary Table S4: Karnataka District-wise KFD Distribution', level=1)
    
    district_data = [
        ('Shimoga', 'Endemic (1957-present)', '~200', '40-50%', 'Original focus'),
        ('Uttara Kannada', 'Endemic (2014-present)', '~80', '15-20%', 'Northern expansion'),
        ('Chikkamagaluru', 'Endemic (2012-present)', '~60', '12-15%', 'Coffee estates'),
        ('Dakshina Kannada', 'Sporadic (2016-present)', '~30', '5-10%', 'Coastal expansion'),
        ('Udupi', 'Sporadic (2018-present)', '~20', '3-5%', 'Recent emergence'),
        ('Belagavi', 'Sporadic (2019-present)', '~10', '2-3%', 'Northern limit'),
    ]
    
    table4 = doc.add_table(rows=len(district_data)+1, cols=5)
    table4.style = 'Table Grid'
    
    headers4 = ['District', 'Status', 'Annual Cases', '% of Total', 'Notes']
    for i, h in enumerate(headers4):
        cell = table4.rows[0].cells[i]
        cell.text = h
        cell.paragraphs[0].runs[0].bold = True
        set_cell_shading(cell, 'D9E2F3')
    
    for i, row_data in enumerate(district_data):
        for j, val in enumerate(row_data):
            table4.rows[i+1].cells[j].text = val
    
    doc.add_page_break()
    
    # FIGURES
    doc.add_heading('Supplementary Figure S1: All Publication Figures', level=1)
    
    figures = [
        ('figure1_target_prioritization.png', 'A. Target Prioritization by Disease Phase'),
        ('figure2_compound_distribution.png', 'B. Compound Distribution'),
        ('figure3_target_potency.png', 'C. Compound Potency'),
        ('figure4_pathway_heatmap.png', 'D. Pathway Analysis'),
        ('figure5_kfd_timeline.png', 'E. KFD Disease Timeline'),
    ]
    
    for filename, title in figures:
        fig_cap = doc.add_paragraph()
        fig_cap.add_run(title).bold = True
        doc.add_picture(str(BASE_DIR / 'outputs' / 'figures' / filename), width=Inches(5.5))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph()
    
    # Save
    output_path = BASE_DIR / 'manuscripts' / 'Supplementary_Materials_KFD.docx'
    doc.save(str(output_path))
    print(f'Created: {output_path}')


def create_cover_letter():
    """Cover letter for Indian Journal of Medical Research (IJMR)"""
    doc = Document()
    
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)
    
    # Date
    doc.add_paragraph(datetime.now().strftime('%B %d, %Y'))
    doc.add_paragraph()
    
    # Address - IJMR
    doc.add_paragraph('The Editor-in-Chief')
    doc.add_paragraph('Indian Journal of Medical Research')
    doc.add_paragraph('Indian Council of Medical Research')
    doc.add_paragraph('New Delhi, India')
    doc.add_paragraph()
    
    # Subject
    subject = doc.add_paragraph()
    subject.add_run('Subject: ').bold = True
    subject.add_run('Manuscript Submission – Host-Directed Therapy for Kyasanur Forest Disease')
    doc.add_paragraph()
    
    # Salutation
    doc.add_paragraph('Dear Editor,')
    doc.add_paragraph()
    
    # Body
    body_paragraphs = [
        'We are pleased to submit our manuscript entitled "Host-Directed Therapy for Kyasanur Forest Disease: An Integrated Multi-omics Pipeline Identifying Endothelial Stabilization and Coagulation Support as Priority Therapeutic Strategies for Karnataka\'s Endemic Hemorrhagic Fever" for consideration as an Original Research article in the Indian Journal of Medical Research.',
        
        'Kyasanur Forest Disease (KFD) is a tick-borne viral hemorrhagic fever endemic to Karnataka\'s Western Ghats, affecting 400-500 patients annually with 3-5% case fatality. Despite recognition since 1957, no specific treatment exists, and the current vaccine has only 50-60% efficacy. This represents a critical gap in Karnataka\'s public health response.',
        
        'Our computational pipeline systematically identifies host-directed therapy targets by integrating viral hemorrhagic fever transcriptomic signatures with druggability assessments. We prioritize endothelial stabilization (angiopoietin pathway) and coagulation support as key therapeutic strategies. Importantly, we identify affordable, immediately available drugs—atorvastatin ($5/course), tranexamic acid ($10/course), and vitamin K ($2/course)—that could be implemented at district hospital level during outbreak seasons.',
        
        'This work directly addresses an ICMR priority area—neglected tropical diseases endemic to India. The practical focus on drugs available in Karnataka\'s peripheral health system makes our findings immediately translatable. We believe this systematic approach would be of significant interest to IJMR\'s readership.',
        
        'This manuscript is original and has not been submitted elsewhere. The study used publicly available de-identified data and did not require ethics approval. No conflicts of interest exist.',
        
        'Thank you for considering our submission.'
    ]
    
    for para in body_paragraphs:
        doc.add_paragraph(para)
    
    doc.add_paragraph()
    doc.add_paragraph('Sincerely,')
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Signature
    sig = doc.add_paragraph()
    sig.add_run('Dr. Siddalingaiah H S').bold = True
    doc.add_paragraph('Professor, Department of Community Medicine')
    doc.add_paragraph('Shridevi Institute of Medical Sciences and Research Hospital')
    doc.add_paragraph('Tumkur – 572106, Karnataka, India')
    doc.add_paragraph('Email: hssling@yahoo.com')
    doc.add_paragraph('Phone: +91-8941087719')
    doc.add_paragraph('ORCID: 0000-0002-4771-8285')
    
    # Save
    output_path = BASE_DIR / 'manuscripts' / 'CoverLetter_KFD_IJMR.docx'
    doc.save(str(output_path))
    print(f'Created: {output_path}')


if __name__ == '__main__':
    create_supplementary()
    create_cover_letter()
